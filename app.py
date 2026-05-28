"""Streamlit UI for the offline interview analysis pipeline."""

import base64
import csv
import io
import json
import re
import sys
import tempfile
import time
from datetime import datetime
from pathlib import Path

import streamlit as st
import yaml

INSTALL_DIR  = Path(__file__).parent
PIPELINE_DIR = INSTALL_DIR / "pipeline"
ASSETS_DIR   = INSTALL_DIR / "assets"
sys.path.insert(0, str(PIPELINE_DIR))

from anonymise import anonymise_transcript  # noqa: E402
from analyse import summarise, extract_themes, analyse_sentiment  # noqa: E402
from questions import analyse_questions  # noqa: E402
from demographics import extract_demographics  # noqa: E402
from compare import build_corpus_comparison  # noqa: E402
from timing import (  # noqa: E402
    estimate_seconds,
    format_duration,
    percent_complete,
    stage_label,
)
from charts import theme_cooccurrence_chart, theme_frequency_chart  # noqa: E402


# ── RIECS brand CSS ───────────────────────────────────────────────────────────

_RIECS_CSS = """
<style>
:root {
    --navy:   #2c324c;
    --steel:  #648a9e;
    --sage:   #85ab86;
    --teal:   #376782;
    --cream:  #f5f2ea;
    --white:  #ffffff;
    --muted:  #7a8a9a;
    --border: #dde4ea;
}

/* ── Hide native Streamlit chrome; use our own titlebar ── */
header[data-testid="stHeader"]   { display: none !important; }
[data-testid="stSidebar"]        { display: none !important; }
[data-testid="collapsedControl"] { display: none !important; }
.stApp { background-color: var(--cream); }
.main .block-container { padding-top: 4.75rem !important; }

/* ── Custom title bar ── */
.riecs-titlebar {
    position: fixed;
    top: 0; left: 0; right: 0;
    height: 3.5rem;
    z-index: 9999;
    background: var(--navy);
    display: flex;
    align-items: center;
    gap: 0.75rem;
    padding: 0 1.5rem;
    box-shadow: 0 2px 8px rgba(0,0,0,0.25);
}
.riecs-titlebar img         { height: 2rem; width: auto; }
.riecs-titlebar-title       { color: #fff; font-weight: 700; font-size: 1.05rem; line-height: 1.2; }
.riecs-titlebar-sub         { color: #8aa8bc; font-size: 0.7rem; margin-top: 0.05rem; }

/* ── Headings ── */
h1 { color: var(--navy) !important; }
h2 {
    color: var(--navy) !important;
    border-bottom: 2px solid var(--border);
    padding-bottom: 0.4rem;
    margin-bottom: 1rem;
}
h3 { color: var(--navy) !important; }

/* ── Primary button ── */
[data-testid="stBaseButton-primary"], button[kind="primary"] {
    background-color: var(--teal) !important;
    color: var(--white) !important;
    border: none !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
    transition: filter 0.15s;
}
[data-testid="stBaseButton-primary"]:hover, button[kind="primary"]:hover {
    filter: brightness(0.88) !important;
}
[data-testid="stBaseButton-primary"]:disabled,
button[kind="primary"]:disabled { opacity: 0.45 !important; }

/* ── Download button ── */
[data-testid="stDownloadButton"] button {
    background-color: var(--sage) !important;
    color: var(--white) !important;
    border: none !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
}
[data-testid="stDownloadButton"] button:hover { filter: brightness(0.90) !important; }

/* ── Progress bar ── */
[data-testid="stProgressBar"] > div {
    background-color: var(--border) !important;
    border-radius: 99px !important;
    height: 10px !important;
}
[data-testid="stProgressBar"] > div > div {
    background-color: var(--teal) !important;
    border-radius: 99px !important;
    transition: width 0.3s ease;
}

/* ── Tabs ── */
[data-baseweb="tab-list"] {
    background: transparent !important;
    border-bottom: 2px solid var(--border) !important;
    gap: 0 !important;
}
[data-baseweb="tab"] {
    color: var(--muted) !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
    background: transparent !important;
    border-bottom: 3px solid transparent !important;
    padding: 0.5rem 1.25rem !important;
}
[aria-selected="true"][data-baseweb="tab"] {
    color: var(--teal) !important;
    border-bottom: 3px solid var(--teal) !important;
}
[data-baseweb="tab-highlight"] { display: none !important; }

/* ── Expanders ── */
[data-testid="stExpander"] {
    background: var(--white) !important;
    border-radius: 12px !important;
    box-shadow: 0 2px 8px rgba(44,50,76,0.07) !important;
    border: 1px solid var(--border) !important;
    margin-bottom: 0.5rem;
}
[data-testid="stExpander"] summary {
    color: var(--navy) !important;
    font-weight: 600 !important;
}

/* ── Metrics ── */
[data-testid="stMetricValue"] {
    color: var(--navy) !important;
    font-weight: 700 !important;
    font-size: 1.75rem !important;
}
[data-testid="stMetricLabel"] {
    font-size: 0.75rem !important;
    text-transform: uppercase !important;
    letter-spacing: 0.06em !important;
    color: var(--steel) !important;
}

/* ── Dataframe ── */
[data-testid="stDataFrame"] {
    border-radius: 8px;
    overflow: hidden;
    box-shadow: 0 2px 8px rgba(44,50,76,0.07);
}

/* ── Selectbox ── */
[data-testid="stSelectbox"] [data-baseweb="select"] > div:first-child {
    background: var(--white) !important;
    border-color: var(--border) !important;
    border-radius: 6px !important;
}

/* ── File uploader ── */
[data-testid="stFileUploader"] section {
    border: 2px dashed var(--border) !important;
    border-radius: 8px !important;
    background: var(--white) !important;
}
[data-testid="stFileUploader"] section:focus-within,
[data-testid="stFileUploader"] section:hover {
    border-color: var(--teal) !important;
}

/* ── Alert / info banners ── */
[data-testid="stAlert"] {
    border-radius: 8px !important;
    border-left: 4px solid var(--steel) !important;
}

/* ── Caption ── */
[data-testid="stCaptionContainer"] {
    color: var(--muted) !important;
    font-size: 0.8rem !important;
}

/* ── Blockquotes ── */
blockquote {
    border-left: 4px solid var(--border);
    padding: 0.5rem 1rem;
    background: #f0f4f7;
    border-radius: 0 6px 6px 0;
    color: var(--steel);
    margin: 0.5rem 0;
}

/* ── Scrollable report pane (Outcomes tab) ── */
.riecs-scroll-pane {
    max-height: 65vh;
    overflow-y: auto;
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 1.5rem 2rem;
    background: var(--white);
    box-shadow: 0 2px 8px rgba(44,50,76,0.07);
    font-size: 0.92rem;
    line-height: 1.75;
}
.riecs-scroll-pane h2 {
    color: var(--navy);
    border-bottom: 1px solid var(--border);
    padding-bottom: 0.3rem;
    margin: 1.5rem 0 0.75rem;
    font-size: 1.1rem;
}
.riecs-scroll-pane h3 {
    color: var(--teal);
    margin: 1.2rem 0 0.4rem;
    font-size: 0.95rem;
}
.riecs-scroll-pane h4 {
    color: var(--steel);
    text-transform: uppercase;
    letter-spacing: 0.05em;
    font-size: 0.78rem;
    margin: 1rem 0 0.3rem;
}
.riecs-scroll-pane blockquote {
    border-left: 3px solid var(--teal);
    background: #f0f4f7;
    padding: 0.4rem 0.9rem;
    margin: 0.4rem 0;
    color: var(--steel);
    font-style: italic;
    border-radius: 0 4px 4px 0;
}
.riecs-scroll-pane .interview-card {
    border-top: 3px solid var(--teal);
    background: var(--cream);
    border-radius: 0 0 8px 8px;
    padding: 1rem 1.25rem;
    margin: 1.5rem 0;
}
.riecs-scroll-pane .badge {
    display: inline-block; padding: 1px 8px; border-radius: 99px;
    font-size: 0.75rem; font-weight: 600;
}
.tone-positive { background: #d4ead5; color: #1a5c2a; }
.tone-negative { background: #fde8e8; color: #7a2020; }
.tone-neutral  { background: var(--border); color: var(--muted); }
.tone-mixed    { background: #fef6e8; color: #7a4a10; }
.freq-high   { background: #e8f0f5; color: var(--teal); }
.freq-medium { background: #f0f4f7; color: var(--steel); }
.freq-low    { background: #f5f7f9; color: var(--muted); }
</style>
"""

# ── Codebook spreadsheet helpers ─────────────────────────────────────────────

_COL_SYNONYMS = {
    "code":        {"code", "code_name", "code name", "id", "key", "identifier", "theme_code"},
    "label":       {"label", "name", "theme", "theme_name", "theme name", "title", "category"},
    "description": {"description", "definition", "desc", "meaning", "notes", "explanation"},
}


def _auto_detect(headers: list[str], field: str) -> int:
    synonyms = _COL_SYNONYMS[field]
    for i, h in enumerate(headers):
        if h.lower().strip() in synonyms:
            return i
    return 0


def parse_spreadsheet(file_bytes: bytes, filename: str) -> tuple[list[dict], list[str]]:
    ext = Path(filename).suffix.lower()
    if ext in (".xlsx", ".xls"):
        import openpyxl
        wb = openpyxl.load_workbook(
            filename=io.BytesIO(file_bytes), read_only=True, data_only=True
        )
        ws = wb.active
        all_rows = list(ws.iter_rows(values_only=True))
        if not all_rows:
            return [], []
        headers = [str(h).strip() if h is not None else f"col_{i}"
                   for i, h in enumerate(all_rows[0])]
        data = [
            {h: (str(v).strip() if v is not None else "") for h, v in zip(headers, row)}
            for row in all_rows[1:]
            if any(v not in (None, "") for v in row)
        ]
        return data, headers
    else:
        text = file_bytes.decode("utf-8-sig")
        reader = csv.DictReader(io.StringIO(text))
        data = [{k: (v or "").strip() for k, v in row.items()} for row in reader]
        headers = list(reader.fieldnames or [])
        return data, headers


# ── Question-list helpers (workshop mode) ────────────────────────────────────

# Strip optional leading "Q1:", "Q1.", "1)", "1.", etc., from a question line.
_QUESTION_PREFIX_RE = re.compile(r"^\s*(?:[Qq])?\s*\d+\s*[\.\):\-]\s*")
# Inline kind markers, e.g. "[aggregate] How many participants in total?"
_KIND_MARKER_RE = re.compile(r"^\s*\[(?P<kind>aggregate|per_document)\]\s*", re.IGNORECASE)


def _read_question_lines(file_bytes: bytes, filename: str) -> list[str]:
    """Extract non-empty question lines from a .txt or .docx upload."""
    ext = Path(filename).suffix.lower()
    if ext == ".docx":
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        raw_lines = [p.text for p in doc.paragraphs]
    else:
        text = file_bytes.decode("utf-8-sig", errors="replace")
        raw_lines = text.splitlines()
    lines = []
    for ln in raw_lines:
        # Preserve any leading kind marker, then strip the numeric prefix.
        kind_match = _KIND_MARKER_RE.match(ln)
        prefix = ""
        body = ln
        if kind_match:
            prefix = kind_match.group(0)
            body = ln[kind_match.end():]
        stripped = prefix + _QUESTION_PREFIX_RE.sub("", body).strip()
        stripped = stripped.strip()
        if stripped and not stripped.startswith("#"):
            lines.append(stripped)
    return lines


def parse_questions_files(files) -> list[dict]:
    """Concatenate one-question-per-line files in upload order, assign q01… ids.

    Each entry is `{id, text, kind}` where `kind` is `per_document` (default)
    or `aggregate` (corpus-level question that's answered from aggregated data
    rather than per-document analysis — flag with an inline `[aggregate]`
    marker, e.g. `[aggregate] How many participants in total?`).

    De-duplication by case-insensitive whitespace-collapsed text — first
    occurrence wins.
    """
    seen: set[str] = set()
    out: list[dict] = []
    for f in files or []:
        for line in _read_question_lines(f.read(), f.name):
            kind = "per_document"
            m = _KIND_MARKER_RE.match(line)
            if m:
                kind = m.group("kind").lower()
                line = line[m.end():].strip()
            if not line:
                continue
            key = re.sub(r"\s+", " ", line).lower()
            if key in seen:
                continue
            seen.add(key)
            out.append({
                "id":   f"q{len(out) + 1:02d}",
                "text": line,
                "kind": kind,
            })
        try:
            f.seek(0)
        except Exception:
            pass
    return out


# Tokens that should not contribute to filename-similarity scoring.
_DEDUP_STOPWORDS = {
    "workshop", "workshops", "ws", "session", "sessions", "template", "templates",
    "draft", "draught", "final", "version", "ver", "rev", "copy", "of", "the", "and",
    "part", "p", "a", "b", "c", "d", "i", "ii", "iii", "iv",
    "notes", "note", "report", "sheet", "description", "observations",
    "2023", "2024", "2025", "2026", "2027",
}


def _workshop_id_for_index(i: int) -> str:
    """Stable, zero-padded workshop identifier."""
    return f"workshop_{i + 1:02d}"


def _assign_workshop_ids(filenames: list[str]) -> dict[str, str]:
    """Assign workshop_01, workshop_02, … to each filename in upload order.

    Returns a mapping `file_stem -> workshop_id`. The mapping is one-per-file
    by design; potential duplicates are reported separately so the researcher
    can verify them before merging anything manually.
    """
    return {
        Path(fname).stem: _workshop_id_for_index(i)
        for i, fname in enumerate(filenames)
    }


def _normalise_for_dedup(name: str) -> set[str]:
    """Tokenise a filename for similarity scoring.

    Lowercase the stem, replace common separators with spaces, drop
    boilerplate tokens (workshop / draft / version-markers / common years /
    single-letter part markers), then return the surviving tokens as a set.
    Pure-digit tokens are preserved — they're typically the workshop number
    that discriminates one session from another.
    """
    stem = Path(name).stem.lower()
    cleaned = re.sub(r"[_\-.()\[\]+&,]+", " ", stem)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    keep: list[str] = []
    for t in cleaned.split():
        if not t:
            continue
        if t in _DEDUP_STOPWORDS:
            continue
        # Explicit version markers ("v2", "ver3", "rev1")
        if re.fullmatch(r"(?:v|ver|rev)\d+[a-z]?", t):
            continue
        # Single-character alpha tokens are part markers (a, b, c …); skip.
        if len(t) == 1 and not t.isdigit():
            continue
        keep.append(t)
    return set(keep)


def _jaccard(a: set, b: set) -> float:
    if not a or not b:
        return 0.0
    return len(a & b) / len(a | b)


def detect_workshop_duplicates(
    filenames: list[str],
    workshop_ids: dict[str, str],
    *,
    threshold: float = 0.7,
) -> list[dict]:
    """Cluster filenames by normalised-token Jaccard similarity.

    Returns a list of clusters, each containing two or more files that may
    describe the same workshop:

        [{"members": [{"file": "...", "workshop_id": "workshop_03"}, ...],
          "min_similarity": 0.83},
         ...]

    The result is empty when no candidate clusters are found.
    """
    if len(filenames) < 2:
        return []

    norm = {f: _normalise_for_dedup(f) for f in filenames}
    parent = {f: f for f in filenames}

    def find(x: str) -> str:
        while parent[x] != x:
            parent[x] = parent[parent[x]]
            x = parent[x]
        return x

    pair_score: dict[tuple[str, str], float] = {}
    for i, a in enumerate(filenames):
        for b in filenames[i + 1:]:
            s = _jaccard(norm[a], norm[b])
            if s >= threshold:
                ra, rb = find(a), find(b)
                if ra != rb:
                    parent[ra] = rb
                pair_score[(a, b)] = s

    from collections import defaultdict as _dd
    clusters: dict[str, list[str]] = _dd(list)
    for f in filenames:
        clusters[find(f)].append(f)

    out: list[dict] = []
    for members in clusters.values():
        if len(members) < 2:
            continue
        # Within-cluster similarity range — the LLM uses this to weigh the hint.
        scores = [
            pair_score[(a, b)]
            for i, a in enumerate(members) for b in members[i + 1:]
            if (a, b) in pair_score
        ]
        out.append({
            "members": [
                {
                    "file": fn,
                    "workshop_id": workshop_ids.get(Path(fn).stem, ""),
                }
                for fn in sorted(members)
            ],
            "min_similarity": round(min(scores), 3) if scores else 0.0,
            "max_similarity": round(max(scores), 3) if scores else 0.0,
        })
    # Surface the strongest-looking clusters first.
    out.sort(key=lambda c: (-c["min_similarity"], -len(c["members"])))
    return out


def codebook_rows_to_yaml(
    rows: list[dict],
    code_col: str,
    label_col: str,
    desc_col: str,
) -> str:
    entries = []
    for row in rows:
        entry: dict = {}
        if row.get(code_col):
            entry["code"] = row[code_col]
        if row.get(label_col):
            entry["label"] = row[label_col]
        if row.get(desc_col):
            entry["description"] = row[desc_col]
        for k, v in row.items():
            if k not in (code_col, label_col, desc_col) and v:
                entry[k] = v
        if entry:
            entries.append(entry)
    return yaml.dump(entries, allow_unicode=True, default_flow_style=False)


# ── Transcript reader ─────────────────────────────────────────────────────────

def read_transcript(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(path)
        lines = []
        for block in doc.element.body:
            tag = block.tag.split("}")[-1]
            if tag == "p":
                text = "".join(t.text or "" for t in block.iter() if t.tag.endswith("}t"))
                if text.strip():
                    lines.append(text)
            elif tag == "tbl":
                for tr in block.iter():
                    if tr.tag.endswith("}tr"):
                        seen: set[int] = set()
                        cells = []
                        for tc in tr.iter():
                            if tc.tag.endswith("}tc") and id(tc) not in seen:
                                seen.add(id(tc))
                                cell_text = "".join(
                                    t.text or "" for t in tc.iter() if t.tag.endswith("}t")
                                )
                                if cell_text.strip():
                                    cells.append(cell_text.strip())
                        if cells:
                            lines.append(" | ".join(cells))
        return "\n\n".join(lines)
    else:
        return path.read_text(encoding="utf-8")


# ── Config ────────────────────────────────────────────────────────────────────

def load_cfg(
    codebook_path: Path | None = None,
    *,
    mode: str | None = None,
    questions_path: Path | None = None,
    anonymise_workshop: bool | None = None,
) -> dict:
    # The installer copies config.yaml to the project root; running from the
    # dev repo it stays under pipeline/. Honour either location.
    _cfg_candidates = [INSTALL_DIR / "config.yaml", PIPELINE_DIR / "config.yaml"]
    _cfg_file = next((p for p in _cfg_candidates if p.exists()), _cfg_candidates[0])
    cfg = yaml.safe_load(_cfg_file.read_text())
    if mode:
        cfg["mode"] = mode
    if codebook_path:
        cfg.setdefault("paths", {})["codebook"] = str(codebook_path)
    if questions_path:
        cfg.setdefault("paths", {})["questions"] = str(questions_path)
    if anonymise_workshop is not None:
        cfg.setdefault("analysis", {})["anonymise_workshop_sheets"] = bool(anonymise_workshop)
    return cfg


# ── Per-interview pipeline helpers ────────────────────────────────────────────

def _process_one_interview(
    interview_path: Path,
    run_dir: Path,
    cfg: dict,
    on_stage_start,
    on_stage_tick,
    on_stage_done,
) -> dict:
    """Run the full per-document pipeline.

    Two modes:
      interviews  — anonymise → summarise → themes → sentiment
      workshop    — [anonymise if cfg.analysis.anonymise_workshop_sheets]
                    → summarise → questions   (sentiment is per question,
                    returned inside the `questions` stage output)
    """
    iid          = interview_path.stem
    entities_dir = run_dir / cfg["gdpr"]["entities_subdir"]
    mode         = (cfg.get("mode") or "interviews").lower()
    is_workshop  = mode == "workshop"
    anonymise_enabled = (
        cfg.get("analysis", {}).get("anonymise_workshop_sheets", True)
        if is_workshop else True
    )
    result: dict = {}

    raw_text   = read_transcript(interview_path)
    word_count = len(raw_text.split())
    timings: dict = {}

    def run_stage(stage: str, fn):
        estimate = estimate_seconds(stage, word_count)
        on_stage_start(stage, estimate)
        t0 = time.time()

        def tick(tokens, elapsed, note=""):
            on_stage_tick(stage, elapsed, estimate, tokens, note)

        out = fn(tick)
        duration = time.time() - t0
        timings[stage] = duration
        on_stage_done(stage, duration)
        return out

    # Stage 1 — anonymise (skippable in workshop mode)
    if anonymise_enabled:
        anon_text, entity_map = run_stage(
            "anonymise",
            lambda tick: anonymise_transcript(raw_text, cfg, tick_cb=tick),
        )
        (run_dir / "anonymised" / f"{iid}_anon.txt").write_text(anon_text, encoding="utf-8")
        (entities_dir / f"{iid}_entities.json").write_text(
            json.dumps(entity_map, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    else:
        anon_text = raw_text  # workshop mode with anonymisation disabled

    # Stage 2 — summarise (both modes)
    summary = run_stage(
        "summarise",
        lambda tick: summarise(anon_text, iid, cfg, tick_cb=tick),
    )
    (run_dir / "analysis" / f"{iid}_summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    result["summary"] = summary

    if is_workshop:
        # Stage 3 (workshop) — per-question coverage / sentiment / emerging themes
        findings = run_stage(
            "questions",
            lambda tick: analyse_questions(anon_text, iid, cfg, tick_cb=tick),
        )
        (run_dir / "analysis" / f"{iid}_questions.json").write_text(
            json.dumps(findings, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        result["questions"] = findings

        # Stage 3b (workshop) — structured demographic facts
        demographics = run_stage(
            "demographics",
            lambda tick: extract_demographics(anon_text, iid, cfg, tick_cb=tick),
        )
        (run_dir / "analysis" / f"{iid}_demographics.json").write_text(
            json.dumps(demographics, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        result["demographics"] = demographics
    else:
        # Stage 3 (interviews) — thematic coding
        themes = run_stage(
            "themes",
            lambda tick: extract_themes(anon_text, iid, cfg, tick_cb=tick),
        )
        (run_dir / "analysis" / f"{iid}_themes.json").write_text(
            json.dumps(themes, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        result["themes"] = themes

        # Stage 4 (interviews only) — overall sentiment
        sentiment = run_stage(
            "sentiment",
            lambda tick: analyse_sentiment(anon_text, iid, cfg, tick_cb=tick),
        )
        (run_dir / "analysis" / f"{iid}_sentiment.json").write_text(
            json.dumps(sentiment, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        result["sentiment"] = sentiment

    (run_dir / "analysis" / f"{iid}_timings.json").write_text(
        json.dumps(timings, indent=2), encoding="utf-8"
    )
    result["_timings"] = timings
    return result


def _load_interview_results(iid: str, run_dir: Path) -> dict:
    """Re-hydrate a previously-completed document's results from disk.

    Loads whichever per-stage JSON files exist; works for both interview and
    workshop runs (extra `_questions.json` for the latter).
    """
    result: dict = {}
    for key, suffix in [
        ("summary",      "_summary"),
        ("themes",       "_themes"),
        ("sentiment",    "_sentiment"),
        ("questions",    "_questions"),
        ("demographics", "_demographics"),
    ]:
        f = run_dir / "analysis" / f"{iid}{suffix}.json"
        if f.exists():
            result[key] = json.loads(f.read_text(encoding="utf-8"))
    timings_f = run_dir / "analysis" / f"{iid}_timings.json"
    if timings_f.exists():
        result["_timings"] = json.loads(timings_f.read_text(encoding="utf-8"))
    return result


def _write_checkpoint(
    run_dir: Path,
    work_dir: Path,
    all_interviews: list[str],
    codebook_path,
    completed: list[str],
    *,
    mode: str = "interviews",
    questions_path=None,
    anonymise_workshop: bool | None = None,
) -> None:
    cp = {
        "run_dir":            str(run_dir),
        "work_dir":           str(work_dir),
        "interviews":         all_interviews,
        "codebook_path":      str(codebook_path) if codebook_path else None,
        "questions_path":     str(questions_path) if questions_path else None,
        "mode":               mode,
        "anonymise_workshop": anonymise_workshop,
        "completed":          completed,
    }
    (run_dir / "checkpoint.json").write_text(json.dumps(cp, indent=2), encoding="utf-8")


def _list_loadable_runs() -> list[Path]:
    """Return run directories with a finished corpus comparison, newest first."""
    output_dir = INSTALL_DIR / "output"
    if not output_dir.exists():
        return []
    runs = []
    for report in output_dir.glob("*/corpus/comparison_report.md"):
        runs.append(report.parent.parent)
    return sorted(runs, key=lambda p: p.name, reverse=True)


def _load_complete_run(run_dir: Path) -> dict:
    """Re-hydrate a finished run's results dict from disk.

    Picks workshop vs interviews mode based on which matrix file is
    present, then walks each per-document JSON in `analysis/` and merges in
    the corpus matrix + report + potential_duplicates + workshop_ids map
    (when the run is from a workshop session).
    """
    analysis_dir = run_dir / "analysis"
    corpus_dir   = run_dir / "corpus"

    questions_matrix = corpus_dir / "questions_matrix.json"
    themes_matrix    = corpus_dir / "themes_matrix.json"
    is_workshop = questions_matrix.exists() and not themes_matrix.exists()
    mode = "workshop" if is_workshop else "interviews"

    # Per-document results — discover IDs from the summary files.
    results: dict = {}
    for sf in sorted(analysis_dir.glob("*_summary.json")):
        iid = sf.stem.replace("_summary", "")
        results[iid] = _load_interview_results(iid, run_dir)

    # Corpus block
    corpus: dict = {}
    matrix_file = questions_matrix if is_workshop else themes_matrix
    if matrix_file.exists():
        corpus["matrix"] = json.loads(matrix_file.read_text(encoding="utf-8"))
    report_file = corpus_dir / "comparison_report.md"
    if report_file.exists():
        corpus["report"] = report_file.read_text(encoding="utf-8")
    dupes_file = corpus_dir / "potential_duplicates.json"
    if dupes_file.exists():
        corpus["_potential_duplicates"] = json.loads(dupes_file.read_text(encoding="utf-8"))

    results["_corpus"] = corpus
    results["_mode"]   = mode

    if is_workshop:
        ws_file = run_dir / "workshops.yaml"
        if ws_file.exists():
            results["_workshop_ids"] = yaml.safe_load(ws_file.read_text(encoding="utf-8")) or {}
        else:
            results["_workshop_ids"] = {}

    return results


def _find_resumable_run() -> dict | None:
    output_dir = INSTALL_DIR / "output"
    if not output_dir.exists():
        return None
    for cp_file in sorted(output_dir.glob("*/checkpoint.json"), reverse=True):
        try:
            cp         = json.loads(cp_file.read_text(encoding="utf-8"))
            completed  = cp.get("completed", [])
            interviews = cp.get("interviews", [])
            if completed and len(completed) < len(interviews) and Path(cp["work_dir"]).exists():
                return cp
        except Exception:
            continue
    return None


# ── HTML report ───────────────────────────────────────────────────────────────

def _logo_data_uri() -> str:
    p = ASSETS_DIR / "riecs-glyph.png"
    if p.exists():
        return "data:image/png;base64," + base64.b64encode(p.read_bytes()).decode()
    return ""


def _md_to_html(text: str) -> str:
    text = re.sub(r"^### (.+)$", r"<h4>\1</h4>", text, flags=re.MULTILINE)
    text = re.sub(r"^## (.+)$",  r"<h3>\1</h3>",  text, flags=re.MULTILINE)
    text = re.sub(r"^# (.+)$",   r"<h2>\1</h2>",   text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"\*(.+?)\*",     r"<em>\1</em>",         text)
    lines = text.split("\n")
    out, in_list = [], False
    for line in lines:
        if re.match(r"^- .+", line):
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(re.sub(r"^- (.+)", r"<li>\1</li>", line))
        else:
            if in_list:
                out.append("</ul>")
                in_list = False
            out.append(line)
    if in_list:
        out.append("</ul>")
    text = "\n".join(out)
    text = re.sub(r"\n{2,}", "</p><p>", text)
    return f"<p>{text}</p>"


_REPORT_CSS = """
:root{--navy:#2c324c;--steel:#648a9e;--sage:#85ab86;--teal:#376782;
      --cream:#f5f2ea;--white:#fff;--muted:#7a8a9a;--border:#dde4ea}
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:system-ui,-apple-system,sans-serif;background:var(--cream);
     color:var(--navy);line-height:1.6;min-height:100vh}
.page{max-width:960px;margin:0 auto;padding:2rem 1.5rem}
.report-header{display:flex;align-items:center;gap:1rem;background:var(--navy);
  color:#fff;padding:1rem 1.5rem;border-radius:12px;margin-bottom:1.5rem}
.report-header img{height:3rem;width:auto}
.report-header h1{font-size:1.3rem;color:#fff;margin:0}
.report-header .sub{font-size:0.78rem;text-transform:uppercase;
  letter-spacing:0.08em;color:#648a9e;margin-top:0.15rem}
.report-meta{font-size:0.78rem;color:var(--muted);margin-bottom:2rem}
h2{color:var(--navy);border-bottom:2px solid var(--border);
   padding-bottom:0.4rem;margin:2rem 0 1rem}
h3{color:var(--navy);margin:1.5rem 0 0.5rem}
h4{color:var(--steel);margin:1rem 0 0.4rem;font-size:0.85rem;
   text-transform:uppercase;letter-spacing:0.05em}
.card{background:var(--white);border-radius:12px;padding:1.5rem;
  box-shadow:0 2px 8px rgba(44,50,76,0.07);margin-bottom:1.5rem;
  border-top:4px solid var(--teal)}
blockquote{margin:.5rem 0;padding:.5rem 1rem;
  border-left:4px solid var(--border);background:#f0f4f7;
  border-radius:0 6px 6px 0;color:var(--steel);font-size:.9rem}
ul{padding-left:1.5rem;margin:.4rem 0}li{margin:.25rem 0}
table{border-collapse:collapse;width:100%;margin:.75rem 0;font-size:.88rem;
  background:var(--white);border-radius:8px;overflow:hidden;
  box-shadow:0 2px 8px rgba(44,50,76,0.07)}
th{background:#f0f4f7;text-align:left;padding:.55rem .75rem;
   border:1px solid var(--border);font-size:.78rem;
   text-transform:uppercase;letter-spacing:.05em;color:var(--steel)}
td{padding:.5rem .75rem;border:1px solid var(--border)}
.badge{display:inline-block;padding:2px 10px;border-radius:99px;
  font-size:.78rem;font-weight:600}
.tone-positive{background:#d4ead5;color:#1a5c2a}
.tone-negative{background:#fde8e8;color:#7a2020}
.tone-neutral{background:var(--border);color:var(--muted)}
.tone-mixed{background:#fef6e8;color:#7a4a10}
.freq-high{background:#e8f0f5;color:var(--teal)}
.freq-medium{background:#f0f4f7;color:var(--steel)}
.freq-low{background:#f5f7f9;color:var(--muted)}
code{background:#f0f4f7;padding:1px 5px;border-radius:3px;font-size:.85em}
.meta{color:var(--muted);font-size:.8rem}
a{color:var(--teal);text-decoration:none}
"""


# ── Demographics aggregation + rendering (workshop mode) ─────────────────────

_AGE_BUCKETS_ORDER = [
    "under_18", "18_24", "25_34", "35_44", "45_54", "55_64", "65_plus", "unspecified",
]
_AGE_BUCKETS_LABEL = {
    "under_18": "Under 18",
    "18_24":    "18–24",
    "25_34":    "25–34",
    "35_44":    "35–44",
    "45_54":    "45–54",
    "55_64":    "55–64",
    "65_plus":  "65+",
    "unspecified": "Unspecified",
}
_GENDER_ORDER = ["female", "male", "non_binary", "unspecified"]
_GENDER_LABEL = {
    "female":      "Female",
    "male":        "Male",
    "non_binary":  "Non-binary",
    "unspecified": "Unspecified",
}
_MODALITY_ORDER = ["on_site", "online", "hybrid", "unspecified"]
_MODALITY_LABEL = {
    "on_site":     "On-site",
    "online":      "Online",
    "hybrid":      "Hybrid",
    "unspecified": "Unspecified",
}


def _normalise_group_label(label: str) -> str:
    """Lightly normalise a stakeholder-group label for cross-doc aggregation.

    Lowercases, trims whitespace, collapses internal spacing, and strips
    a trailing 's' so singular/plural variants merge. This is intentionally
    conservative — anything semantic is left to the researcher (via the
    taxonomy file consulted by _resolve_stakeholder_label).
    """
    s = (label or "").strip().lower()
    s = re.sub(r"\s+", " ", s)
    if len(s) > 3 and s.endswith("s") and not s.endswith("ss"):
        s = s[:-1]
    return s


def _load_stakeholder_taxonomy() -> dict[str, str]:
    """Read stakeholder_taxonomy.yaml and return alias_key -> canonical label.

    Looks at INSTALL_DIR/stakeholder_taxonomy.yaml first (production layout)
    and project root. Missing file → empty mapping (fallback behaviour is
    the heuristic in _normalise_group_label).
    """
    candidates = [
        INSTALL_DIR / "stakeholder_taxonomy.yaml",
        INSTALL_DIR.parent / "stakeholder_taxonomy.yaml",
    ]
    path = next((p for p in candidates if p.exists()), None)
    if path is None:
        return {}
    try:
        data = yaml.safe_load(path.read_text(encoding="utf-8")) or []
    except Exception:
        return {}
    mapping: dict[str, str] = {}
    for entry in data:
        if not isinstance(entry, dict):
            continue
        canonical = (entry.get("canonical") or "").strip()
        if not canonical:
            continue
        # The canonical label is itself an alias for itself.
        mapping[_normalise_group_label(canonical)] = canonical
        for a in entry.get("aliases") or []:
            # If a user wrote `- foo: bar` unquoted, YAML parses it as a dict
            # rather than a string. Salvage both halves so we don't silently
            # drop the alias.
            if isinstance(a, dict):
                for k, v in a.items():
                    candidate = f"{k}: {v}" if v not in (None, "") else str(k)
                    key = _normalise_group_label(candidate)
                    if key:
                        mapping[key] = canonical
                continue
            key = _normalise_group_label(str(a))
            if key:
                mapping[key] = canonical
    return mapping


def _resolve_stakeholder_label(label: str, taxonomy: dict[str, str]) -> tuple[str, str]:
    """Map a raw stakeholder label to (canonical_label, normalised_key).

    Falls back to the lightly-normalised label when no taxonomy entry matches.
    """
    norm = _normalise_group_label(label)
    if norm in taxonomy:
        canonical = taxonomy[norm]
        return canonical, _normalise_group_label(canonical)
    return label.strip(), norm


def _aggregate_demographics(interviews: dict, workshop_ids: dict) -> dict:
    """Roll per-doc demographics into corpus-wide totals.

    Per-workshop rows preserve the original order from the interviews dict
    (which mirrors processing order). Stakeholder labels are first resolved
    against `stakeholder_taxonomy.yaml` (so e.g. `Policy Makers` and
    `policymakers` both become `Policymakers & government`); anything not in
    the taxonomy falls back to the lightly-normalised raw label.
    """
    rows: list[dict] = []
    grand_total = 0
    gender_total = {g: 0 for g in _GENDER_ORDER}
    age_total    = {b: 0 for b in _AGE_BUCKETS_ORDER}
    modality_total = {m: 0 for m in _MODALITY_ORDER}
    group_totals_n: dict[str, int] = {}
    # normalised-key -> first-seen display label.
    group_first_label: dict[str, str] = {}
    # Track which raw labels got merged where, for the caveats section.
    raw_labels_seen: set[str] = set()
    taxonomy_merged: dict[str, set[str]] = {}  # canonical -> set of raw labels

    taxonomy = _load_stakeholder_taxonomy()

    for doc_id, data in interviews.items():
        demo = data.get("demographics") or {}
        if not demo:
            continue
        wid = (workshop_ids or {}).get(doc_id, "")
        n_p = demo.get("n_participants") or 0
        gender_raw = demo.get("gender") or {}
        gender = {g: (gender_raw.get(g) or 0) for g in _GENDER_ORDER}
        # If n_participants is missing but gender counts are present, derive it.
        if not n_p and sum(gender.values()) > 0:
            n_p = sum(gender.values())
        modality = demo.get("modality") or "unspecified"
        if modality not in _MODALITY_ORDER:
            modality = "unspecified"
        ages_raw = {
            b["bucket"]: int(b.get("n") or 0)
            for b in (demo.get("age_buckets") or [])
            if b.get("bucket") in _AGE_BUCKETS_ORDER
        }
        groups = demo.get("stakeholder_groups") or []

        rows.append({
            "doc_id":      doc_id,
            "workshop_id": wid,
            "n":           n_p,
            "gender":      gender,
            "modality":    modality,
            "ages":        ages_raw,
            "groups":      groups,
            "notes":       demo.get("notes"),
        })

        grand_total += n_p
        for g, c in gender.items():
            gender_total[g] += c
        for b, c in ages_raw.items():
            age_total[b] = age_total.get(b, 0) + c
        modality_total[modality] += 1
        for gr in groups:
            raw_label = gr.get("group", "")
            n         = int(gr.get("n") or 0)
            if not raw_label:
                continue
            raw_labels_seen.add(raw_label.strip())
            display, key = _resolve_stakeholder_label(raw_label, taxonomy)
            if not key:
                continue
            group_totals_n[key] = group_totals_n.get(key, 0) + n
            group_first_label.setdefault(key, display)
            taxonomy_merged.setdefault(display, set()).add(raw_label.strip())

    group_rows = sorted(
        [{"label": group_first_label[k], "n": v}
         for k, v in group_totals_n.items()],
        key=lambda r: (-r["n"], r["label"].lower()),
    )

    return {
        "rows":             rows,
        "grand_total":      grand_total,
        "gender_total":     gender_total,
        "age_total":        age_total,
        "modality_total":   modality_total,
        "groups":           group_rows,
        "n_documents":      sum(1 for r in rows),
        "raw_label_count":  len(raw_labels_seen),
        "taxonomy_applied": bool(taxonomy),
        "taxonomy_merges":  {
            canonical: sorted(raws)
            for canonical, raws in taxonomy_merged.items()
            if len(raws) > 1
        },
    }


def _age_pyramid_chart(age_total: dict, gender_total: dict, *, title: str = "Age distribution") -> bytes | None:
    """Render an age distribution bar chart.

    A true left/right pyramid needs per-gender age counts — which we don't
    extract today (the prompt asks for totals, not the cross). For now this
    function renders a horizontal-bar age histogram in brand colours. The
    function name keeps the contract clean so we can swap in a real pyramid
    later without touching the renderers.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import io as _io

    buckets = [b for b in _AGE_BUCKETS_ORDER if (age_total.get(b) or 0) > 0]
    if not buckets:
        return None
    labels = [_AGE_BUCKETS_LABEL[b] for b in buckets]
    counts = [age_total[b] for b in buckets]

    fig, ax = plt.subplots(figsize=(7, max(2.5, 0.55 * len(buckets) + 1.0)))
    bars = ax.barh(labels, counts, color="#376782", height=0.65)
    for bar, count in zip(bars, counts):
        ax.text(bar.get_width() + max(counts) * 0.02,
                bar.get_y() + bar.get_height() / 2,
                str(count), va="center", ha="left", fontsize=9, color="#2c324c")
    ax.set_xlim(0, max(counts) * 1.18 if counts else 1)
    ax.set_xlabel("Participants", fontsize=9, color="#2c324c")
    ax.set_title(title, fontsize=12, color="#2c324c", fontweight="bold", pad=10)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    ax.set_axisbelow(True)
    ax.xaxis.grid(True, color="#e4e4e4", linewidth=0.8)
    ax.invert_yaxis()  # youngest at top
    buf = _io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return buf.getvalue()


def _compute_caveats(results: dict, demo_agg: dict | None) -> list[str]:
    """Build the data-driven Caveats bullets for a workshop run.

    Each item is a short, plain-English sentence anchored in numbers from the
    run. The list is intentionally specific — generic disclaimers belong in
    the methodology section, not here.
    """
    bullets: list[str] = []
    interviews = {k: v for k, v in results.items() if not k.startswith("_")}
    n_files     = len(interviews)
    corpus      = results.get("_corpus") or {}
    matrix      = (corpus.get("matrix") or {}).get("codes") or {}
    n_questions = len(matrix)
    dupes       = corpus.get("_potential_duplicates") or []
    workshop_ids = results.get("_workshop_ids") or {}
    n_workshops = len({workshop_ids.get(k, k) for k in interviews})

    # Files vs workshops
    if n_workshops != n_files:
        bullets.append(
            f"{n_files} source files cover {n_workshops} distinct workshops "
            "(file-to-workshop mapping in the Files Evaluated table)."
        )
    else:
        bullets.append(
            f"Each of the {n_files} source files is treated as one workshop. "
            "Files that may actually describe the same workshop are listed in the "
            "Potential Duplicate Workshops section above and need researcher review."
        )

    # Workshop IDs
    if workshop_ids:
        bullets.append(
            "workshop_NN identifiers were assigned in upload order, not "
            "chronological order; cross-reference the date in the filename if needed."
        )

    # Demographics coverage (only meaningful when extraction ran)
    if demo_agg and demo_agg.get("rows"):
        total_participants = demo_agg["grand_total"]
        n_demo_docs        = demo_agg["n_documents"]
        gender_total       = sum(demo_agg["gender_total"].values())
        age_total          = sum(demo_agg["age_total"].values())
        no_count           = sum(1 for r in demo_agg["rows"] if not r["n"])

        bullets.append(
            f"Demographics were extracted from {n_demo_docs} of {n_files} "
            "workshop description sheets."
        )
        if no_count:
            bullets.append(
                f"{no_count} workshop(s) did not state a participant count in the "
                "description; their `n_participants` is reported as 0."
            )
        if total_participants:
            bullets.append(
                f"Gender was reported for {gender_total} of {total_participants} "
                f"participants ({100 * gender_total / total_participants:.0f}%); the rest "
                "are absent because the workshop sheet did not state a breakdown."
            )
            bullets.append(
                f"Age was reported for {age_total} of {total_participants} "
                f"participants ({100 * age_total / total_participants:.0f}%); the age "
                "distribution chart shows only the subset where ages were stated."
            )

        # Stakeholder taxonomy merges
        if demo_agg.get("taxonomy_applied"):
            n_raw       = demo_agg.get("raw_label_count") or 0
            n_canonical = len(demo_agg.get("groups") or [])
            merges      = demo_agg.get("taxonomy_merges") or {}
            bullets.append(
                f"Stakeholder labels were collapsed from {n_raw} raw variants to "
                f"{n_canonical} canonical groups via stakeholder_taxonomy.yaml "
                f"({len(merges)} canonical group(s) merged two or more aliases)."
            )

    # Potential duplicates
    if dupes:
        bullets.append(
            f"{len(dupes)} cluster(s) of filename-similar workshops were flagged as "
            "possible duplicates; the executive summary discusses each candidate. "
            "Workshop counts above treat every file as a distinct workshop unless "
            "you merge them manually."
        )

    # Models used (read from config; the prompt-stage caveat matters because
    # llama3.1:8b is known to hallucinate counts on long structured prompts).
    try:
        models = (load_cfg().get("models") or {})
        compare_model = models.get("compare", "(unset)")
        bullets.append(
            f"Per-document analysis (anonymise / summarise / questions / "
            f"demographics) used `{models.get('questions', '(unset)')}`. "
            f"Cross-workshop synthesis used `{compare_model}` — chosen because the "
            "smaller model frequently confuses matrix row counts with workshop counts."
        )
    except Exception:
        pass

    # Question coverage gap
    if n_questions and matrix:
        not_answered_total = 0
        partially_total    = 0
        for info in matrix.values():
            by_iv = info.get("by_interview") or {}
            partially_total    += sum(1 for v in by_iv.values() if v == "partially_answered")
            # not_answered is implicit when a doc id is absent from by_interview
        cells_filled = sum(len(info.get("by_interview") or {}) for info in matrix.values())
        cells_total  = n_files * n_questions
        if cells_total and cells_filled < cells_total:
            absent = cells_total - cells_filled
            bullets.append(
                f"{absent} of {cells_total} workshop×question cells are flagged "
                f"`not answered` (the document did not address that question)."
            )

    return bullets


def _caveats_html(bullets: list[str]) -> str:
    if not bullets:
        return ""
    items = "".join(f"<li>{b}</li>" for b in bullets)
    return (
        '<h2>Caveats</h2>'
        '<p class="meta">How to read the rest of this report. These notes are '
        'computed from this specific run and are not generic disclaimers.</p>'
        f'<ul>{items}</ul>'
    )


def _demographics_html(demo_agg: dict, age_png_b64: str | None) -> str:
    """Render the demographics section into HTML."""
    if not demo_agg or not demo_agg.get("rows"):
        return ""

    rows         = demo_agg["rows"]
    grand_total  = demo_agg["grand_total"]
    gender_total = demo_agg["gender_total"]
    modality_t   = demo_agg["modality_total"]
    groups       = demo_agg["groups"]
    n_documents  = demo_agg["n_documents"]

    parts: list[str] = ['<h2>Demographics</h2>']

    # ── Participants & modality table ────────────────────────────────────────
    parts.append(
        '<h3>Participants &amp; modality</h3>'
        '<table><tr><th>Workshop</th>'
        '<th style="text-align:right">Participants</th><th>Modality</th></tr>'
    )
    for r in rows:
        parts.append(
            f'<tr><td><code>{r["workshop_id"]}</code></td>'
            f'<td style="text-align:right">{r["n"] or "&mdash;"}</td>'
            f'<td>{_MODALITY_LABEL.get(r["modality"], r["modality"])}</td></tr>'
        )
    parts.append(
        f'<tr style="font-weight:700;background:#f0f4f7">'
        f'<td>Total ({n_documents} workshop(s))</td>'
        f'<td style="text-align:right">{grand_total}</td>'
        f'<td>'
        + ", ".join(
            f'{_MODALITY_LABEL[m]}: {modality_t.get(m,0)}'
            for m in _MODALITY_ORDER if modality_t.get(m, 0) > 0
        )
        + '</td></tr></table>'
    )

    # ── Gender distribution table (percent + count) ─────────────────────────
    parts.append(
        '<h3>Gender distribution</h3>'
        '<table><tr><th>Workshop</th>'
        + "".join(f'<th style="text-align:right">{_GENDER_LABEL[g]} %</th>' for g in _GENDER_ORDER)
        + '<th style="text-align:right">N</th></tr>'
    )
    def _pct(part: int, total: int) -> str:
        return f"{100 * part / total:.0f}%" if total else "&mdash;"
    for r in rows:
        total = sum(r["gender"].values()) or (r["n"] or 0)
        parts.append(
            f'<tr><td><code>{r["workshop_id"]}</code></td>'
            + "".join(
                f'<td style="text-align:right">{_pct(r["gender"][g], total)}</td>'
                for g in _GENDER_ORDER
            )
            + f'<td style="text-align:right">{total}</td></tr>'
        )
    total_known_gender = sum(gender_total.values())
    parts.append(
        f'<tr style="font-weight:700;background:#f0f4f7">'
        f'<td>Total</td>'
        + "".join(
            f'<td style="text-align:right">{_pct(gender_total[g], total_known_gender)}</td>'
            for g in _GENDER_ORDER
        )
        + f'<td style="text-align:right">{total_known_gender}</td></tr></table>'
    )

    # ── Stakeholder groups ───────────────────────────────────────────────────
    if groups:
        parts.append(
            '<h3>Stakeholder groups (total across workshops)</h3>'
            '<table><tr><th>Group</th>'
            '<th style="text-align:right">Participants</th>'
            '<th style="text-align:right">Share</th></tr>'
        )
        tot_g = sum(g["n"] for g in groups) or 1
        for g in groups:
            parts.append(
                f'<tr><td>{g["label"]}</td>'
                f'<td style="text-align:right">{g["n"]}</td>'
                f'<td style="text-align:right">{100 * g["n"] / tot_g:.0f}%</td></tr>'
            )
        parts.append('</table>')

    # ── Age distribution chart ───────────────────────────────────────────────
    if age_png_b64:
        parts.append(
            '<h3>Age distribution</h3>'
            f'<img src="data:image/png;base64,{age_png_b64}" '
            f'style="max-width:100%;height:auto" alt="Age distribution">'
        )

    return "\n".join(parts)


def _interview_stats(
    interviews: dict,
    mode: str = "interviews",
    workshop_ids: dict | None = None,
) -> list[dict]:
    """Per-document stats for the report's 'Files Evaluated' section.

    Interviews mode counts themes; workshop mode counts questions the
    document actually addressed (coverage != "not_answered") and surfaces the
    workshop_id assigned to each file.
    """
    workshop_ids = workshop_ids or {}
    stats = []
    for iid, data in interviews.items():
        summary = data.get("summary", {}) or {}
        timings = data.get("_timings", {}) or {}
        if mode == "workshop":
            qs = (data.get("questions", {}) or {}).get("questions", []) or []
            n_main = sum(1 for q in qs if (q.get("coverage") or "") != "not_answered")
        else:
            themes = (data.get("themes", {}) or {}).get("themes", []) or []
            n_main = len(themes)
        stats.append({
            "id":          iid,
            "workshop_id": workshop_ids.get(iid, ""),
            "words":       summary.get("word_count"),
            "n_main":      n_main,
            "seconds":     sum(timings.values()) if timings else None,
        })
    return stats


# Coverage strings → CSS badge class shared by HTML report + scroll pane.
_COVERAGE_CLASS = {
    "answered":           "freq-high",
    "partially_answered": "freq-medium",
    "not_answered":       "freq-low",
}

# Traffic-light palette for the transposed coverage matrix (workshop mode).
# `bg` / `fg` are also used by the DOCX table via cell shading.
_COVERAGE_STYLE = {
    "answered":           {"bg": "#4caf50", "fg": "#ffffff",
                           "label": "answered",           "symbol": "✓"},
    "partially_answered": {"bg": "#ffc107", "fg": "#3a3a3a",
                           "label": "partially answered", "symbol": "½"},
    "not_answered":       {"bg": "#e53935", "fg": "#ffffff",
                           "label": "not answered",       "symbol": "—"},
}
_COVERAGE_ORDER = ["answered", "partially_answered", "not_answered"]


def _docs_in_workshop_order(
    matrix_codes: dict,
    workshop_ids: dict,
) -> list[tuple[str, str]]:
    """Return (workshop_id, document_id) pairs in workshop_id order.

    Falls back to alphabetic doc-id ordering when no workshop_ids map exists.
    Documents that appear in the matrix but have no workshop_id (e.g. older
    runs) get a `—` placeholder so they still render.
    """
    # Every doc id we know about — from the workshop_ids map plus any extras
    # mentioned in the matrix.
    all_doc_ids: set[str] = set(workshop_ids or {})
    for info in matrix_codes.values():
        all_doc_ids.update((info.get("by_interview") or {}).keys())

    def _key(d: str) -> tuple[str, str]:
        wid = (workshop_ids or {}).get(d, "")
        return (wid or "￿", d)  # docs without a wid sink to the bottom

    return [((workshop_ids or {}).get(d) or "—", d)
            for d in sorted(all_doc_ids, key=_key)]


def _coverage_for(matrix_codes: dict, qid: str, doc_id: str) -> str:
    """Look up a (question, document) cell. Missing → 'not_answered'."""
    cov = (matrix_codes.get(qid, {}).get("by_interview") or {}).get(doc_id)
    return cov if cov in _COVERAGE_STYLE else "not_answered"


def _workshop_coverage_html(
    matrix_codes: dict,
    workshop_ids: dict,
) -> str:
    """Transposed, colour-coded coverage table for the HTML report.

    Rows = workshops (sorted by workshop_id), columns = q01..qNN. Cells are
    coloured cold→hot (not_answered → partially_answered → answered) and
    carry a glyph for accessibility. Followed by a colour legend and a
    `q## → question text` key.
    """
    if not matrix_codes:
        return ""

    qids = list(matrix_codes.keys())
    doc_pairs = _docs_in_workshop_order(matrix_codes, workshop_ids)

    header = (
        '<tr><th style="text-align:left">Workshop</th>'
        + "".join(
            f'<th style="text-align:center;padding:0.4rem 0.3rem"><code>{q}</code></th>'
            for q in qids
        )
        + "</tr>"
    )

    rows: list[str] = []
    for wid, doc_id in doc_pairs:
        cells = [
            f'<td style="font-weight:600;white-space:nowrap;padding:0.35rem 0.6rem">'
            f'<code>{wid}</code></td>'
        ]
        for qid in qids:
            cov   = _coverage_for(matrix_codes, qid, doc_id)
            style = _COVERAGE_STYLE[cov]
            cells.append(
                f'<td title="{style["label"]}: {doc_id} / {qid}" '
                f'style="background:{style["bg"]};color:{style["fg"]};'
                f'text-align:center;font-weight:700;padding:0.3rem 0.4rem">'
                f'{style["symbol"]}</td>'
            )
        rows.append("<tr>" + "".join(cells) + "</tr>")

    # Legend
    legend_items = "".join(
        f'<span style="display:inline-block;width:14px;height:14px;'
        f'background:{_COVERAGE_STYLE[c]["bg"]};border:1px solid #c5c8cf;'
        f'border-radius:3px;vertical-align:middle;margin-right:0.35rem"></span>'
        f'<span style="margin-right:1.2rem">{_COVERAGE_STYLE[c]["symbol"]} '
        f'{_COVERAGE_STYLE[c]["label"]}</span>'
        for c in _COVERAGE_ORDER
    )

    # Question key
    q_key = "".join(
        f'<li><code>{qid}</code> &mdash; {matrix_codes[qid].get("label", qid)}</li>'
        for qid in qids
    )

    return (
        '<table style="border-collapse:separate;border-spacing:2px;'
        'font-size:0.85rem;width:auto">'
        f'{header}{"".join(rows)}'
        '</table>'
        f'<p class="meta" style="margin-top:0.6rem"><strong>Legend:</strong> {legend_items}</p>'
        f'<h4 style="margin-top:0.8rem">Question key</h4>'
        f'<ul style="font-size:0.85rem">{q_key}</ul>'
    )


def _mode_strings(mode: str) -> dict:
    """Centralised mode-aware label strings used across the three report
    renderers, so wording stays consistent."""
    if mode == "workshop":
        return {
            "title":             "Workshop Templates Analysis Report",
            "doc_singular":      "Workshop document",
            "doc_plural":        "workshop document(s)",
            "files_count_col":   "Questions answered",
            "matrix_heading":    "Question Coverage Matrix",
            "main_col":          "Question",
            "main_id_col":       "ID",
            "freq_chart_title":  "Question coverage",
            "freq_chart_xlabel": "Coverage score  (answered = 3, partial = 2, not answered = 0, "
                                 "summed across documents)",
            "freq_chart_noun":   "document",
            "cooc_chart_title":  "Question co-occurrence  (documents covering both)",
            "section_heading":   "Questions",
        }
    return {
        "title":             "Interview Analysis Report",
        "doc_singular":      "Interview",
        "doc_plural":        "interview(s)",
        "files_count_col":   "Themes",
        "matrix_heading":    "Theme Matrix",
        "main_col":          "Theme",
        "main_id_col":       "Code",
        "freq_chart_title":  "Theme relevance",
        "freq_chart_xlabel": "Relevance score  (high = 3, medium = 2, low = 1, "
                             "summed across interviews)",
        "freq_chart_noun":   "interview",
        "cooc_chart_title":  "Theme co-occurrence  (shared interviews)",
        "section_heading":   "Themes",
    }


def generate_html_report(run_dir: Path, results: dict) -> str:
    now      = datetime.now().strftime("%Y-%m-%d %H:%M")
    corpus   = results.get("_corpus", {})
    mode     = results.get("_mode") or "interviews"
    is_workshop = mode == "workshop"
    L        = _mode_strings(mode)
    workshop_ids = results.get("_workshop_ids") or {}
    interviews = {k: v for k, v in results.items() if not k.startswith("_")}
    logo_uri = _logo_data_uri()
    logo_img = f'<img src="{logo_uri}" alt="RIECS">' if logo_uri else ""

    # Count distinct workshops when in workshop mode; for interviews this is
    # just the file count.
    if is_workshop and workshop_ids:
        _n_workshops = len({workshop_ids.get(k, k) for k in interviews})
        _doc_count_str = (
            f'{len(interviews)} file(s) &mdash; {_n_workshops} workshop(s)'
        )
    else:
        _doc_count_str = f'{len(interviews)} {L["doc_plural"]}'

    parts = [
        f'<!DOCTYPE html><html lang="en"><head><meta charset="utf-8">'
        f'<title>{L["title"]} — {now}</title>'
        f'<style>{_REPORT_CSS}</style></head><body><div class="page">'
        f'<div class="report-header">{logo_img}'
        f'<div><h1>{L["title"]}</h1>'
        f'<div class="sub">RIECS — offline analysis</div></div></div>'
        f'<p class="report-meta">Generated {now} &mdash; {_doc_count_str}'
        f' &mdash; output: {run_dir}</p>'
    ]

    stats = _interview_stats(interviews, mode=mode, workshop_ids=workshop_ids)
    if stats:
        rows = ""
        for s in stats:
            secs = format_duration(s["seconds"]) if s["seconds"] else "&mdash;"
            ws_cell = f'<td><code>{s["workshop_id"]}</code></td>' if is_workshop else ""
            rows += (
                f'<tr>{ws_cell}<td>{s["id"]}</td>'
                f'<td>{s["words"] if s["words"] is not None else "&mdash;"}</td>'
                f'<td>{s["n_main"]}</td><td>{secs}</td></tr>'
            )
        if corpus.get("_duration"):
            empty_first = '<td>&mdash;</td>' if is_workshop else ""
            rows += (
                f'<tr>{empty_first}<td><em>Corpus comparison</em></td><td>&mdash;</td>'
                f'<td>&mdash;</td><td>{format_duration(corpus["_duration"])}</td></tr>'
            )
        ws_header = '<th>Workshop</th>' if is_workshop else ""
        parts.append(
            f'<h2>Files Evaluated</h2><table>'
            f'<tr>{ws_header}<th>{L["doc_singular"]}</th><th>Words</th><th>{L["files_count_col"]}</th>'
            f'<th>Processing time</th></tr>{rows}</table>'
        )

    # Heads-up about potentially duplicated workshops, drawn from the
    # heuristic clusters detected before corpus comparison.
    _dupes = corpus.get("_potential_duplicates") or []
    if is_workshop and _dupes:
        items = ""
        for cl in _dupes:
            members = ", ".join(
                f'<code>{m.get("workshop_id","")}</code> {m["file"]}'
                for m in cl["members"]
            )
            items += (
                f'<li>{members}'
                f' &nbsp; <span class="meta">similarity '
                f'{cl["min_similarity"]:.2f}–{cl["max_similarity"]:.2f}</span></li>'
            )
        parts.append(
            f'<h2>Potential Duplicate Workshops</h2>'
            f'<p class="meta">Filename similarity suggests the following files may describe '
            f'the same workshop. The executive summary discusses each candidate; '
            f'verify before merging.</p>'
            f'<ul>{items}</ul>'
        )

    if corpus.get("report"):
        parts.append(f'<h2>Executive Summary</h2>{_md_to_html(corpus["report"])}')

    # Caveats — computed from this run, positioned right after the executive
    # summary so readers see them before interpreting the deeper tables.
    if is_workshop:
        _demo_for_caveats = _aggregate_demographics(interviews, workshop_ids)
        parts.append(_caveats_html(_compute_caveats(results, _demo_for_caveats)))

    matrix_codes = corpus.get("matrix", {}).get("codes", {})
    if matrix_codes:
        parts.append(f'<h2>{L["matrix_heading"]}</h2>')
        if is_workshop:
            parts.append(_workshop_coverage_html(matrix_codes, workshop_ids))
        else:
            iids   = sorted({iid for c in matrix_codes.values() for iid in c["by_interview"]})
            header = (
                f"<tr><th>{L['main_col']}</th><th>{L['main_id_col']}</th>"
                + "".join(f"<th>{i}</th>" for i in iids)
                + f"<th>{L['doc_plural'].replace('(s)', 's').title()}</th></tr>"
            )
            rows = "".join(
                f'<tr><td>{info["label"]}</td><td><code>{code}</code></td>'
                + "".join(f'<td>{info["by_interview"].get(i,"&mdash;")}</td>' for i in iids)
                + f'<td>{info["total_interviews"]}</td></tr>'
                for code, info in matrix_codes.items()
            )
            parts.append(f'<table>{header}{rows}</table>')

    if matrix_codes:
        # Workshop mode: label nodes/bars by qid (q01…) rather than the full
        # question text — the prose lives in the Question key under the matrix.
        chart_label_map = {c: c for c in matrix_codes} if is_workshop else None
        freq_png = theme_frequency_chart(
            matrix_codes,
            title=L["freq_chart_title"],
            xlabel=L["freq_chart_xlabel"],
            item_noun=L["freq_chart_noun"],
            label_map=chart_label_map,
        )
        if freq_png:
            b64 = base64.b64encode(freq_png).decode()
            parts.append(
                f'<h2>{L["freq_chart_title"].title()}</h2>'
                f'<img src="data:image/png;base64,{b64}" '
                f'style="max-width:100%;height:auto" alt="{L["freq_chart_title"]} chart">'
            )
        cooc_png = theme_cooccurrence_chart(
            matrix_codes,
            title=L["cooc_chart_title"],
            label_map=chart_label_map,
        )
        if cooc_png:
            b64 = base64.b64encode(cooc_png).decode()
            parts.append(
                f'<h2>{L["cooc_chart_title"].split("  ")[0].title()}</h2>'
                f'<img src="data:image/png;base64,{b64}" '
                f'style="max-width:100%;height:auto" alt="{L["cooc_chart_title"]}">'
            )

    # Demographics section — workshop mode only, only when extraction ran.
    if is_workshop:
        demo_agg = _aggregate_demographics(interviews, workshop_ids)
        if demo_agg.get("rows"):
            age_png = _age_pyramid_chart(demo_agg["age_total"], demo_agg["gender_total"])
            age_png_b64 = base64.b64encode(age_png).decode() if age_png else None
            parts.append(_demographics_html(demo_agg, age_png_b64))

    for iid, data in interviews.items():
        _ws = workshop_ids.get(iid, "")
        _hdr_iid = (
            f'<code style="font-size:0.7em;color:#648a9e">{_ws}</code> &nbsp; {iid}'
            if is_workshop and _ws else iid
        )
        parts.append(
            f'<div class="card">'
            f'<h2 style="border:none;margin-top:0">{L["doc_singular"]}: {_hdr_iid}</h2>'
        )
        s = data.get("summary", {})
        if s:
            parts.append(
                f'<h3>Summary</h3>'
                f'<p><strong>Estimated duration:</strong> {s.get("estimated_duration_min","&mdash;")} min'
                f' &nbsp;|&nbsp; <strong>Word count:</strong> {s.get("word_count","&mdash;")}</p>'
            )
            if s.get("key_topics"):
                items = "".join(
                    f'<li><strong>{t["topic"]}</strong>: {t.get("brief_description","")}</li>'
                    if isinstance(t, dict) else f"<li>{t}</li>"
                    for t in s["key_topics"]
                )
                parts.append(f'<h4>Key topics</h4><ul>{items}</ul>')
            if s.get("main_positions"):
                items = "".join(f"<li>{p}</li>" for p in s["main_positions"])
                parts.append(f'<h4>Main positions</h4><ul>{items}</ul>')
            if s.get("notable_quotes"):
                quotes = "".join(f"<blockquote>{q}</blockquote>" for q in s["notable_quotes"])
                parts.append(f'<h4>Notable quotes</h4>{quotes}')
            if s.get("methodological_notes"):
                parts.append(f'<p class="meta"><em>{s["methodological_notes"]}</em></p>')

        if is_workshop:
            q = data.get("questions", {}) or {}
            tone = q.get("overall_tone")
            register = q.get("emotional_register")
            if tone or register:
                parts.append(
                    f'<h3>Document tone</h3>'
                    f'<p>'
                    + (f'Overall tone: <span class="badge tone-{tone}">{tone}</span>' if tone else "")
                    + (f' &nbsp; Register: {register}' if register else "")
                    + '</p>'
                )

            if q.get("questions"):
                parts.append(f'<h3>{L["section_heading"]}</h3>')
                for qf in q["questions"]:
                    coverage = qf.get("coverage", "not_answered")
                    cls      = _COVERAGE_CLASS.get(coverage, "freq-low")
                    sentiment = qf.get("sentiment", "neutral")
                    quotes = "".join(
                        f"<blockquote>{esc}</blockquote>"
                        for esc in qf.get("supporting_quotes", [])[:3]
                    )
                    emerging = qf.get("emerging_themes") or []
                    emerging_html = (
                        f'<p class="meta"><em>Emerging themes:</em> {", ".join(emerging)}</p>'
                        if emerging else ""
                    )
                    answer = qf.get("answer", "") or ""
                    parts.append(
                        f'<h4>{qf.get("question_text", "")} '
                        f'<code>{qf.get("question_id","")}</code> '
                        f'<span class="badge {cls}">{coverage.replace("_", " ")}</span> '
                        f'<span class="badge tone-{sentiment}">{sentiment}</span></h4>'
                        f'<p>{answer}</p>{quotes}{emerging_html}'
                    )
                if q.get("cross_question_notes"):
                    parts.append(
                        f'<p class="meta"><em>Cross-question notes:</em> '
                        f'{q["cross_question_notes"]}</p>'
                    )

        else:
            t = data.get("themes", {})
            if t.get("themes"):
                parts.append(f'<h3>{L["section_heading"]}</h3>')
                for theme in t["themes"]:
                    freq  = theme.get("frequency", "")
                    sub   = (f'<p class="meta">Sub-themes: {", ".join(theme["sub_themes"])}</p>'
                             if theme.get("sub_themes") else "")
                    quotes = "".join(
                        f"<blockquote>{q}</blockquote>"
                        for q in theme.get("supporting_quotes", [])[:3]
                    )
                    parts.append(
                        f'<h4>{theme.get("label", theme.get("code",""))}'
                        f' <code>{theme.get("code","")}</code>'
                        f' <span class="badge freq-{freq}">{freq}</span></h4>'
                        f'<p>{theme.get("description","")}</p>{sub}{quotes}'
                    )
                if t.get("new_codes_proposed"):
                    parts.append(
                        f'<p class="meta"><em>New codes proposed: '
                        f'{", ".join(t["new_codes_proposed"])}</em></p>'
                    )

            sent = data.get("sentiment", {})
            if sent:
                tone = sent.get("overall_tone", "neutral")
                parts.append(
                    f'<h3>Sentiment</h3>'
                    f'<p>Overall tone: <span class="badge tone-{tone}">{tone}</span>'
                    f' &nbsp; Confidence: {sent.get("confidence","&mdash;")}'
                    f' &nbsp; Register: {sent.get("emotional_register","&mdash;")}</p>'
                )
                if sent.get("topic_sentiments"):
                    ts_rows = "".join(
                        f'<tr><td>{ts.get("topic","")}</td>'
                        f'<td><span class="badge tone-{ts.get("tone","neutral")}">'
                        f'{ts.get("tone","")}</span></td>'
                        f'<td>{ts.get("notes","")}</td></tr>'
                        for ts in sent["topic_sentiments"]
                    )
                    parts.append(
                        f'<table><tr><th>Topic</th><th>Tone</th><th>Notes</th></tr>'
                        f'{ts_rows}</table>'
                    )
                if sent.get("notable_passages"):
                    passages = "".join(
                        f"<blockquote>{p}</blockquote>" for p in sent["notable_passages"]
                    )
                    parts.append(f'<h4>Notable passages</h4>{passages}')

        parts.append("</div>")

    parts.append("</div></body></html>")
    return "\n".join(parts)


# ── DOCX report ───────────────────────────────────────────────────────────────

def _docx_shade_cell(cell, hex_color: str) -> None:
    """Apply background shading to a python-docx table cell."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _docx_set_cell_text_color(cell, hex_color: str) -> None:
    """Force a foreground colour on every run in a cell."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    rgb = hex_color.lstrip("#")
    for p in cell.paragraphs:
        for r in p.runs:
            rPr = r._r.get_or_add_rPr()
            col = OxmlElement("w:color")
            col.set(qn("w:val"), rgb)
            rPr.append(col)


def _docx_render_workshop_coverage(doc, matrix_codes: dict, workshop_ids: dict) -> None:
    """Render the transposed colour-coded coverage matrix into a python-docx Document."""
    qids = list(matrix_codes.keys())
    doc_pairs = _docs_in_workshop_order(matrix_codes, workshop_ids)
    n_cols = 1 + len(qids)

    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Workshop"
    for i, qid in enumerate(qids):
        hdr[1 + i].text = qid
    for c in hdr:
        for run in c.paragraphs[0].runs:
            run.bold = True

    for wid, doc_id in doc_pairs:
        row = tbl.add_row().cells
        row[0].text = wid
        for run in row[0].paragraphs[0].runs:
            run.bold = True
        for i, qid in enumerate(qids):
            cov = _coverage_for(matrix_codes, qid, doc_id)
            style = _COVERAGE_STYLE[cov]
            cell = row[1 + i]
            cell.text = style["symbol"]
            _docx_shade_cell(cell, style["bg"])
            _docx_set_cell_text_color(cell, style["fg"])
            for p in cell.paragraphs:
                p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True

    # Legend
    legend = doc.add_paragraph()
    legend.add_run("Legend: ").bold = True
    for c in _COVERAGE_ORDER:
        legend.add_run(f"  {_COVERAGE_STYLE[c]['symbol']} {_COVERAGE_STYLE[c]['label']}    ")

    # Question key
    doc.add_heading("Question key", 3)
    for qid in qids:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{qid}: "); run.bold = True
        p.add_run(matrix_codes[qid].get("label", qid))


def _docx_render_demographics(doc, demo_agg: dict) -> None:
    """Render the Demographics section into a python-docx Document."""
    from docx.shared import Inches

    rows         = demo_agg["rows"]
    grand_total  = demo_agg["grand_total"]
    gender_total = demo_agg["gender_total"]
    modality_t   = demo_agg["modality_total"]
    groups       = demo_agg["groups"]
    n_documents  = demo_agg["n_documents"]

    doc.add_heading("Demographics", 1)

    # Participants & modality
    doc.add_heading("Participants & modality", 2)
    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Table Grid"
    for i, h in enumerate(["Workshop", "Participants", "Modality"]):
        tbl.rows[0].cells[i].text = h
        for run in tbl.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for r in rows:
        row = tbl.add_row().cells
        row[0].text = r["workshop_id"] or ""
        row[1].text = str(r["n"] or "—")
        row[2].text = _MODALITY_LABEL.get(r["modality"], r["modality"])
    total_row = tbl.add_row().cells
    total_row[0].text = f"Total ({n_documents})"
    total_row[1].text = str(grand_total)
    total_row[2].text = ", ".join(
        f'{_MODALITY_LABEL[m]}: {modality_t.get(m,0)}'
        for m in _MODALITY_ORDER if modality_t.get(m, 0) > 0
    )
    for c in total_row:
        for run in c.paragraphs[0].runs:
            run.bold = True

    # Gender distribution
    doc.add_heading("Gender distribution", 2)
    cols = ["Workshop"] + [f"{_GENDER_LABEL[g]} %" for g in _GENDER_ORDER] + ["N"]
    tbl = doc.add_table(rows=1, cols=len(cols)); tbl.style = "Table Grid"
    for i, h in enumerate(cols):
        tbl.rows[0].cells[i].text = h
        for run in tbl.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True

    def _pct(part: int, total: int) -> str:
        return f"{100 * part / total:.0f}%" if total else "—"

    for r in rows:
        total = sum(r["gender"].values()) or (r["n"] or 0)
        row = tbl.add_row().cells
        row[0].text = r["workshop_id"] or ""
        for j, g in enumerate(_GENDER_ORDER):
            row[1 + j].text = _pct(r["gender"][g], total)
        row[len(_GENDER_ORDER) + 1].text = str(total)
    tot_g = sum(gender_total.values())
    row = tbl.add_row().cells
    row[0].text = "Total"
    for j, g in enumerate(_GENDER_ORDER):
        row[1 + j].text = _pct(gender_total[g], tot_g)
    row[len(_GENDER_ORDER) + 1].text = str(tot_g)
    for c in row:
        for run in c.paragraphs[0].runs:
            run.bold = True

    # Stakeholder groups
    if groups:
        doc.add_heading("Stakeholder groups (total across workshops)", 2)
        tbl = doc.add_table(rows=1, cols=3); tbl.style = "Table Grid"
        for i, h in enumerate(["Group", "Participants", "Share"]):
            tbl.rows[0].cells[i].text = h
            for run in tbl.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        tot = sum(g["n"] for g in groups) or 1
        for g in groups:
            row = tbl.add_row().cells
            row[0].text = g["label"]
            row[1].text = str(g["n"])
            row[2].text = f"{100 * g['n'] / tot:.0f}%"

    # Age distribution chart
    age_png = _age_pyramid_chart(demo_agg["age_total"], demo_agg["gender_total"])
    if age_png:
        doc.add_heading("Age distribution", 2)
        doc.add_picture(io.BytesIO(age_png), width=Inches(6.0))


def generate_docx_report(run_dir: Path, results: dict) -> bytes:
    from docx import Document as DocxDoc
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDoc()
    corpus     = results.get("_corpus", {})
    mode       = results.get("_mode") or "interviews"
    is_workshop = mode == "workshop"
    L          = _mode_strings(mode)
    workshop_ids = results.get("_workshop_ids") or {}
    interviews = {k: v for k, v in results.items() if not k.startswith("_")}
    now        = datetime.now().strftime("%Y-%m-%d %H:%M")

    t = doc.add_heading(L["title"], 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {now}").italic = True
    doc.add_paragraph(f"Output: {run_dir}").italic = True

    if is_workshop and workshop_ids:
        _nws = len({workshop_ids.get(k, k) for k in interviews})
        p = doc.add_paragraph(
            f"{len(interviews)} file(s) — {_nws} workshop(s)"
        )
        p.runs[0].italic = True

    stats = _interview_stats(interviews, mode=mode, workshop_ids=workshop_ids)
    if stats:
        doc.add_heading("Files Evaluated", 1)
        ncols = 5 if is_workshop else 4
        tbl = doc.add_table(rows=1, cols=ncols)
        tbl.style = "Table Grid"
        headers = (
            ["Workshop", L["doc_singular"], "Words", L["files_count_col"], "Processing time"]
            if is_workshop else
            [L["doc_singular"], "Words", L["files_count_col"], "Processing time"]
        )
        for i, head in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = head
            for r in cell.paragraphs[0].runs:
                r.bold = True
        for s in stats:
            row = tbl.add_row().cells
            offset = 0
            if is_workshop:
                row[0].text = s["workshop_id"] or ""
                offset = 1
            row[offset + 0].text = s["id"]
            row[offset + 1].text = str(s["words"]) if s["words"] is not None else "—"
            row[offset + 2].text = str(s["n_main"])
            row[offset + 3].text = format_duration(s["seconds"]) if s["seconds"] else "—"
        if corpus.get("_duration"):
            row = tbl.add_row().cells
            offset = 0
            if is_workshop:
                row[0].text = "—"; offset = 1
            row[offset + 0].text = "Corpus comparison"
            row[offset + 1].text = "—"
            row[offset + 2].text = "—"
            row[offset + 3].text = format_duration(corpus["_duration"])

    # Potential duplicates section — researcher-facing flag, before exec summary.
    _dupes = corpus.get("_potential_duplicates") or []
    if is_workshop and _dupes:
        doc.add_heading("Potential Duplicate Workshops", 1)
        p = doc.add_paragraph(
            "Filename similarity suggests the following files may describe "
            "the same workshop. The executive summary discusses each candidate; "
            "verify before merging."
        )
        p.runs[0].italic = True
        for cl in _dupes:
            members = ", ".join(
                f'{m.get("workshop_id","")} {m["file"]}' for m in cl["members"]
            )
            doc.add_paragraph(
                f'{members}  (similarity {cl["min_similarity"]:.2f}–{cl["max_similarity"]:.2f})',
                style="List Bullet",
            )

    if corpus.get("report"):
        doc.add_heading("Executive Summary", 1)
        for line in corpus["report"].split("\n"):
            stripped = line.strip()
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:], 2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], 1)
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped:
                doc.add_paragraph(stripped)

    if is_workshop:
        _demo_for_caveats = _aggregate_demographics(interviews, workshop_ids)
        _caveats = _compute_caveats(results, _demo_for_caveats)
        if _caveats:
            doc.add_heading("Caveats", 1)
            p = doc.add_paragraph(
                "How to read the rest of this report. These notes are computed "
                "from this specific run and are not generic disclaimers."
            )
            p.runs[0].italic = True
            for bullet in _caveats:
                doc.add_paragraph(bullet, style="List Bullet")

    matrix_codes = corpus.get("matrix", {}).get("codes", {})
    if matrix_codes:
        doc.add_heading(L["matrix_heading"], 1)
        if is_workshop:
            _docx_render_workshop_coverage(doc, matrix_codes, workshop_ids)
        else:
            iids  = sorted({iid for c in matrix_codes.values() for iid in c["by_interview"]})
            ncols = 2 + len(iids) + 1
            tbl   = doc.add_table(rows=1, cols=ncols)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text = L["main_col"]
            hdr[1].text = L["main_id_col"]
            for i, iid in enumerate(iids):
                hdr[2 + i].text = iid
            hdr[ncols - 1].text = f"# {L['doc_plural'].replace('(s)', 's').title()}"
            for code, info in matrix_codes.items():
                row = tbl.add_row().cells
                row[0].text = info.get("label", "")
                row[1].text = code
                for i, iid in enumerate(iids):
                    row[2 + i].text = str(info["by_interview"].get(iid, ""))
                row[ncols - 1].text = str(info.get("total_interviews", ""))

    if matrix_codes:
        chart_label_map = {c: c for c in matrix_codes} if is_workshop else None
        freq_png = theme_frequency_chart(
            matrix_codes,
            title=L["freq_chart_title"],
            xlabel=L["freq_chart_xlabel"],
            item_noun=L["freq_chart_noun"],
            label_map=chart_label_map,
        )
        if freq_png:
            doc.add_heading(L["freq_chart_title"].title(), 1)
            doc.add_picture(io.BytesIO(freq_png), width=Inches(6.2))
        cooc_png = theme_cooccurrence_chart(
            matrix_codes,
            title=L["cooc_chart_title"],
            label_map=chart_label_map,
        )
        if cooc_png:
            doc.add_heading(L["cooc_chart_title"].split("  ")[0].title(), 1)
            doc.add_picture(io.BytesIO(cooc_png), width=Inches(6.0))

    if is_workshop:
        demo_agg = _aggregate_demographics(interviews, workshop_ids)
        if demo_agg.get("rows"):
            _docx_render_demographics(doc, demo_agg)

    for iid, data in interviews.items():
        _ws = workshop_ids.get(iid, "")
        _heading = (
            f"{L['doc_singular']}: {_ws} — {iid}"
            if is_workshop and _ws else
            f"{L['doc_singular']}: {iid}"
        )
        doc.add_heading(_heading, 1)

        s = data.get("summary", {})
        if s:
            doc.add_heading("Summary", 2)
            doc.add_paragraph(
                f"Estimated duration: {s.get('estimated_duration_min', '—')} min"
                f"  |  Word count: {s.get('word_count', '—')}"
            )
            if s.get("key_topics"):
                doc.add_heading("Key Topics", 3)
                for topic in s["key_topics"]:
                    if isinstance(topic, dict):
                        p = doc.add_paragraph(style="List Bullet")
                        run = p.add_run(f"{topic['topic']}: ")
                        run.bold = True
                        p.add_run(topic.get("brief_description", ""))
                    else:
                        doc.add_paragraph(str(topic), style="List Bullet")
            if s.get("main_positions"):
                doc.add_heading("Main Positions", 3)
                for pos in s["main_positions"]:
                    doc.add_paragraph(str(pos), style="List Bullet")
            if s.get("notable_quotes"):
                doc.add_heading("Notable Quotes", 3)
                for q in s["notable_quotes"]:
                    p = doc.add_paragraph(f'"{q}"')
                    p.runs[0].italic = True
            if s.get("methodological_notes"):
                p = doc.add_paragraph(s["methodological_notes"])
                p.runs[0].italic = True

        if is_workshop:
            q = data.get("questions", {}) or {}
            tone = q.get("overall_tone")
            register = q.get("emotional_register")
            if tone or register:
                doc.add_heading("Document tone", 2)
                doc.add_paragraph(
                    f"Overall tone: {tone or '—'}  |  Register: {register or '—'}"
                )

            if q.get("questions"):
                doc.add_heading(L["section_heading"], 2)
                for qf in q["questions"]:
                    coverage = (qf.get("coverage") or "not_answered").replace("_", " ")
                    sentiment = qf.get("sentiment", "neutral")
                    heading_text = (
                        f"{qf.get('question_text','')}  "
                        f"[{coverage}  ·  {sentiment}]"
                    )
                    doc.add_heading(heading_text, 3)
                    if qf.get("answer"):
                        doc.add_paragraph(qf["answer"])
                    for quote in qf.get("supporting_quotes", [])[:3]:
                        p = doc.add_paragraph(f'"{quote}"')
                        p.runs[0].italic = True
                    if qf.get("emerging_themes"):
                        p = doc.add_paragraph(
                            "Emerging themes: " + ", ".join(qf["emerging_themes"])
                        )
                        p.runs[0].italic = True
                if q.get("cross_question_notes"):
                    p = doc.add_paragraph(
                        "Cross-question notes: " + str(q["cross_question_notes"])
                    )
                    p.runs[0].italic = True

        else:
            themes = data.get("themes", {})
            if themes.get("themes"):
                doc.add_heading(L["section_heading"], 2)
                for theme in themes["themes"]:
                    label = theme.get("label", theme.get("code", ""))
                    freq  = theme.get("frequency", "")
                    doc.add_heading(f"{label}  [{freq}]", 3)
                    doc.add_paragraph(theme.get("description", ""))
                    if theme.get("sub_themes"):
                        doc.add_paragraph("Sub-themes: " + ", ".join(theme["sub_themes"]))
                    for q in theme.get("supporting_quotes", [])[:3]:
                        p = doc.add_paragraph(f'"{q}"')
                        p.runs[0].italic = True
                if themes.get("new_codes_proposed"):
                    doc.add_paragraph("New codes proposed: " + ", ".join(themes["new_codes_proposed"]))

            sent = data.get("sentiment", {})
            if sent:
                doc.add_heading("Sentiment", 2)
                doc.add_paragraph(
                    f"Overall tone: {sent.get('overall_tone', '—')}"
                    f"  |  Confidence: {sent.get('confidence', '—')}"
                    f"  |  Register: {sent.get('emotional_register', '—')}"
                )
                if sent.get("topic_sentiments"):
                    for ts in sent["topic_sentiments"]:
                        notes = f" — {ts['notes']}" if ts.get("notes") else ""
                        doc.add_paragraph(
                            f"{ts.get('topic', '')}: {ts.get('tone', '')}{notes}",
                            style="List Bullet",
                        )
                if sent.get("notable_passages"):
                    doc.add_heading("Notable Passages", 3)
                    for passage in sent["notable_passages"]:
                        p = doc.add_paragraph(f'"{passage}"')
                        p.runs[0].italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Scrollable pane HTML builder ──────────────────────────────────────────────

def _report_highlights_html(results: dict) -> str:
    corpus     = results.get("_corpus", {})
    mode       = results.get("_mode") or "interviews"
    is_workshop = mode == "workshop"
    L          = _mode_strings(mode)
    workshop_ids = results.get("_workshop_ids") or {}
    interviews = {k: v for k, v in results.items() if not k.startswith("_")}
    parts: list[str] = []

    _dupes = corpus.get("_potential_duplicates") or []
    if is_workshop and _dupes:
        cluster_items = []
        for cl in _dupes:
            members_html = ", ".join(
                f'<code>{m.get("workshop_id", "")}</code> {m["file"]}'
                for m in cl["members"]
            )
            cluster_items.append(
                f'<li>{members_html} &nbsp; <span class="meta">'
                f'similarity {cl["min_similarity"]:.2f}–{cl["max_similarity"]:.2f}</span></li>'
            )
        parts.append(
            f'<h2>Potential Duplicate Workshops</h2>'
            f'<p class="meta">Filename similarity suggests the files below may describe '
            f'the same workshop. The executive summary discusses each candidate; verify before merging.</p>'
            f'<ul>{"".join(cluster_items)}</ul>'
        )

    if corpus.get("report"):
        parts.append(f'<h2>Executive Summary</h2>{_md_to_html(corpus["report"])}')

    # Caveats — computed from this run, positioned right after the executive
    # summary so readers see them before interpreting the deeper tables.
    if is_workshop:
        _demo_for_caveats = _aggregate_demographics(interviews, workshop_ids)
        parts.append(_caveats_html(_compute_caveats(results, _demo_for_caveats)))

    matrix_codes = corpus.get("matrix", {}).get("codes", {})
    if matrix_codes:
        parts.append(f'<h2>{L["matrix_heading"]}</h2>')
        if is_workshop:
            parts.append(_workshop_coverage_html(matrix_codes, workshop_ids))
        else:
            iids   = sorted({iid for c in matrix_codes.values() for iid in c["by_interview"]})
            header = (
                f"<tr><th>{L['main_col']}</th><th>{L['main_id_col']}</th>"
                + "".join(f"<th>{i}</th>" for i in iids)
                + f"<th>{L['doc_plural'].replace('(s)', 's').title()}</th></tr>"
            )
            rows = "".join(
                f'<tr><td>{info["label"]}</td><td><code>{code}</code></td>'
                + "".join(f'<td>{info["by_interview"].get(i, "&mdash;")}</td>' for i in iids)
                + f'<td>{info["total_interviews"]}</td></tr>'
                for code, info in matrix_codes.items()
            )
            parts.append(
                f'<table style="border-collapse:collapse;width:100%;font-size:.85rem">'
                f'<thead style="background:#f0f4f7">{header}</thead>'
                f'<tbody>{rows}</tbody></table>'
            )

    # Charts in the scrollable highlights pane (workshop mode only) — matches
    # the downloadable HTML report so researchers see the same picture in-app.
    if is_workshop and matrix_codes:
        chart_label_map = {c: c for c in matrix_codes}
        freq_png = theme_frequency_chart(
            matrix_codes,
            title=L["freq_chart_title"],
            xlabel=L["freq_chart_xlabel"],
            item_noun=L["freq_chart_noun"],
            label_map=chart_label_map,
        )
        if freq_png:
            b64 = base64.b64encode(freq_png).decode()
            parts.append(
                f'<h2>{L["freq_chart_title"].title()}</h2>'
                f'<img src="data:image/png;base64,{b64}" '
                f'style="max-width:100%;height:auto" alt="{L["freq_chart_title"]}">'
            )
        cooc_png = theme_cooccurrence_chart(
            matrix_codes,
            title=L["cooc_chart_title"],
            label_map=chart_label_map,
        )
        if cooc_png:
            b64 = base64.b64encode(cooc_png).decode()
            parts.append(
                f'<h2>{L["cooc_chart_title"].split("  ")[0].title()}</h2>'
                f'<img src="data:image/png;base64,{b64}" '
                f'style="max-width:100%;height:auto" alt="{L["cooc_chart_title"]}">'
            )

    # Demographics in the scrollable highlights pane (workshop mode only)
    if is_workshop:
        demo_agg = _aggregate_demographics(interviews, workshop_ids)
        if demo_agg.get("rows"):
            age_png = _age_pyramid_chart(demo_agg["age_total"], demo_agg["gender_total"])
            age_b64 = base64.b64encode(age_png).decode() if age_png else None
            parts.append(_demographics_html(demo_agg, age_b64))

    for iid, data in interviews.items():
        _ws = workshop_ids.get(iid, "")
        _h = (
            f'<code style="font-size:0.7em;color:#648a9e">{_ws}</code> &nbsp; {iid}'
            if is_workshop and _ws else iid
        )
        card: list[str] = [f'<div class="interview-card"><h2>{_h}</h2>']

        s = data.get("summary", {})
        if s:
            card.append(
                f'<h3>Summary</h3>'
                f'<p><strong>Duration (est.):</strong> {s.get("estimated_duration_min","—")} min'
                f' &nbsp;|&nbsp; <strong>Words:</strong> {s.get("word_count","—")}</p>'
            )
            if s.get("key_topics"):
                items = "".join(
                    f'<li><strong>{t["topic"]}</strong>: {t.get("brief_description","")}</li>'
                    if isinstance(t, dict) else f"<li>{t}</li>"
                    for t in s["key_topics"]
                )
                card.append(f'<h4>Key topics</h4><ul>{items}</ul>')
            if s.get("main_positions"):
                card.append('<h4>Main positions</h4><ul>'
                            + "".join(f"<li>{p}</li>" for p in s["main_positions"])
                            + "</ul>")
            if s.get("notable_quotes"):
                card.append('<h4>Notable quotes</h4>'
                            + "".join(f"<blockquote>{q}</blockquote>"
                                      for q in s["notable_quotes"]))

        if is_workshop:
            q = data.get("questions", {}) or {}
            if q.get("questions"):
                card.append(f"<h3>{L['section_heading']}</h3>")
                for qf in q["questions"]:
                    coverage = qf.get("coverage", "not_answered")
                    cls      = _COVERAGE_CLASS.get(coverage, "freq-low")
                    sentiment = qf.get("sentiment", "neutral")
                    card.append(
                        f'<h4>{qf.get("question_text", "")} '
                        f'<code>{qf.get("question_id","")}</code> '
                        f'<span class="badge {cls}">{coverage.replace("_", " ")}</span> '
                        f'<span class="badge tone-{sentiment}">{sentiment}</span></h4>'
                        f'<p>{qf.get("answer", "") or ""}</p>'
                    )
                    for quote in qf.get("supporting_quotes", [])[:2]:
                        card.append(f"<blockquote>{quote}</blockquote>")
            tone = q.get("overall_tone")
            register = q.get("emotional_register")
            if tone or register:
                card.append(
                    f'<h3>Document tone</h3>'
                    f'<p>'
                    + (f'Tone: <span class="badge tone-{tone}">{tone}</span>' if tone else "")
                    + (f' &nbsp; Register: {register}' if register else "")
                    + '</p>'
                )
        else:
            themes = data.get("themes", {})
            if themes.get("themes"):
                card.append(f"<h3>{L['section_heading']}</h3>")
                for theme in themes["themes"]:
                    freq = theme.get("frequency", "")
                    card.append(
                        f'<h4>{theme.get("label", theme.get("code",""))}'
                        f' <code>{theme.get("code","")}</code>'
                        f' <span class="badge freq-{freq}">{freq}</span></h4>'
                        f'<p>{theme.get("description","")}</p>'
                    )
                    for q in theme.get("supporting_quotes", [])[:2]:
                        card.append(f"<blockquote>{q}</blockquote>")

            sent = data.get("sentiment", {})
            if sent:
                tone = sent.get("overall_tone", "neutral")
                card.append(
                    f'<h3>Sentiment</h3>'
                    f'<p>Tone: <span class="badge tone-{tone}">{tone}</span>'
                    f' &nbsp; Confidence: {sent.get("confidence","—")}'
                    f' &nbsp; Register: {sent.get("emotional_register","—")}</p>'
                )
                if sent.get("notable_passages"):
                    card.append("".join(
                        f"<blockquote>{p}</blockquote>" for p in sent["notable_passages"]
                    ))

        card.append("</div>")
        parts.append("".join(card))

    return "\n".join(parts)


# ── Streamlit app ─────────────────────────────────────────────────────────────

_favicon = ASSETS_DIR / "favicon.ico"
st.set_page_config(
    page_title="Interview Analyser — RIECS",
    page_icon=str(_favicon) if _favicon.exists() else "📋",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown(_RIECS_CSS, unsafe_allow_html=True)

# Custom title bar
_logo_uri = _logo_data_uri()
_logo_img = f'<img src="{_logo_uri}" alt="RIECS">' if _logo_uri else ""
st.markdown(
    f'<div class="riecs-titlebar">'
    f'{_logo_img}'
    f'<div>'
    f'<div class="riecs-titlebar-title">Interview Analyser</div>'
    f'<div class="riecs-titlebar-sub">Fully offline — all processing happens on this machine</div>'
    f'</div></div>',
    unsafe_allow_html=True,
)

for key, default in [
    ("results", None), ("run_dir", None), ("running", False),
    ("work_dir", None), ("pending", []), ("partial_results", {}),
    ("all_interviews", []), ("codebook_path", None), ("stop_requested", False),
    ("cb_rows", None), ("cb_headers", None),
    ("cb_code_col", None), ("cb_label_col", None), ("cb_desc_col", None),
    # Workshop-mode state
    ("mode", "interviews"),       # "interviews" | "workshop"
    ("questions_list", None),     # parsed list of {id, text}
    ("questions_path", None),     # path to the questions.yaml written into _work
    ("anonymise_workshop", True), # workshop-mode anonymisation toggle
    ("workshop_ids", {}),         # file_stem -> workshop_id (workshop mode)
    ("run_log", []),              # per-doc records persisted across reruns
]:
    if key not in st.session_state:
        st.session_state[key] = default

# ── Mode toggle ───────────────────────────────────────────────────────────────

# Sit above the tabs so it is visible on both Progress and Outcomes.
_MODE_LABELS = {"interviews": "Interviews", "workshop": "Workshop templates"}
_mode_col, _mode_caption = st.columns([1, 3], gap="medium")
with _mode_col:
    _selected_label = st.segmented_control(
        "Analysis mode",
        options=list(_MODE_LABELS.values()),
        default=_MODE_LABELS[st.session_state.mode],
        label_visibility="collapsed",
        disabled=st.session_state.running,
        key="mode_segctrl",
    )
    # Defensive: segmented_control returns None when nothing is selected.
    if _selected_label is None:
        _selected_label = _MODE_LABELS[st.session_state.mode]
    _new_mode = next(k for k, v in _MODE_LABELS.items() if v == _selected_label)
    if _new_mode != st.session_state.mode and not st.session_state.running:
        st.session_state.mode = _new_mode
        st.rerun()
with _mode_caption:
    if st.session_state.mode == "workshop":
        st.caption("Workshop templates — interrogates workshop description sheets against a list of questions.")
    else:
        st.caption("Interviews — thematic coding of interview transcripts against a labelbook.")

# ── Tabs ──────────────────────────────────────────────────────────────────────

tab_progress, tab_outcomes = st.tabs(["Progress", "Outcomes"])

# ── Progress tab ──────────────────────────────────────────────────────────────

with tab_progress:
    col_inputs, col_status = st.columns(2, gap="large")

    # ── Left column: file inputs + run/stop/resume controls ──────────────────

    with col_inputs:
        _is_workshop = st.session_state.mode == "workshop"

        if _is_workshop:
            st.subheader("Workshop description sheets")
            uploaded_txts = st.file_uploader(
                "Workshop documents",
                type=["txt", "docx"],
                accept_multiple_files=True,
                help="Drag and drop plain-text or Word workshop description sheets.",
                label_visibility="collapsed",
                disabled=st.session_state.running,
                key="workshop_uploader",
            )

            st.subheader("Research questions")
            uploaded_questions = st.file_uploader(
                "Question list(s)",
                type=["txt", "docx"],
                accept_multiple_files=True,
                help="One question per line; numeric prefixes (Q1:, 1.) are stripped. "
                     "Multiple files are concatenated in upload order.",
                label_visibility="collapsed",
                disabled=st.session_state.running,
                key="questions_uploader",
            )

            if uploaded_questions and not st.session_state.running:
                _qs = parse_questions_files(uploaded_questions)
                st.session_state.questions_list = _qs or None
                if _qs:
                    with st.expander(f"{len(_qs)} questions parsed — preview", expanded=False):
                        for q in _qs[:8]:
                            st.write(f"- `{q['id']}` — {q['text']}")
                        if len(_qs) > 8:
                            st.caption(f"…and {len(_qs) - 8} more")
                else:
                    st.warning("No questions found in the uploaded file(s).")
            elif not uploaded_questions and not st.session_state.running:
                st.session_state.questions_list = None

            st.session_state.anonymise_workshop = st.checkbox(
                "Anonymise workshop sheets before analysis",
                value=st.session_state.anonymise_workshop,
                help="Recommended. Replaces participant names and other PII with coded tokens.",
                disabled=st.session_state.running,
                key="anonymise_workshop_chk",
            )

            # Workshop mode does not use the labelbook — make sure stale state
            # cannot leak into the run.
            st.session_state.cb_rows    = None
            st.session_state.cb_headers = None

        else:
            st.subheader("Transcripts")
            uploaded_txts = st.file_uploader(
                "Interview files",
                type=["txt", "docx"],
                accept_multiple_files=True,
                help="Drag and drop plain-text or Word transcript files.",
                label_visibility="collapsed",
                disabled=st.session_state.running,
                key="interview_uploader",
            )

            st.subheader("Labelbook")
            uploaded_codebook = st.file_uploader(
                "Labelbook file",
                type=["xlsx", "xls", "csv", "yaml", "yml"],
                help="Excel, CSV, or YAML codebook — optional.",
                label_visibility="collapsed",
                disabled=st.session_state.running,
            )

            if uploaded_codebook and not st.session_state.running:
                ext = Path(uploaded_codebook.name).suffix.lower()
                if ext in (".xlsx", ".xls", ".csv"):
                    file_bytes = uploaded_codebook.read()
                    rows, headers = parse_spreadsheet(file_bytes, uploaded_codebook.name)
                    if rows and headers:
                        if st.session_state.cb_headers != headers:
                            st.session_state.cb_rows      = rows
                            st.session_state.cb_headers   = headers
                            st.session_state.cb_code_col  = headers[_auto_detect(headers, "code")]
                            st.session_state.cb_label_col = headers[_auto_detect(headers, "label")]
                            st.session_state.cb_desc_col  = headers[_auto_detect(headers, "description")]

                        with st.expander(f"{len(rows)} codes — map columns", expanded=False):
                            st.session_state.cb_code_col = st.selectbox(
                                "Code column", headers,
                                index=headers.index(st.session_state.cb_code_col),
                                key="sel_code",
                            )
                            st.session_state.cb_label_col = st.selectbox(
                                "Label column", headers,
                                index=headers.index(st.session_state.cb_label_col),
                                key="sel_label",
                            )
                            st.session_state.cb_desc_col = st.selectbox(
                                "Description column", headers,
                                index=headers.index(st.session_state.cb_desc_col),
                                key="sel_desc",
                            )
                            st.caption("Preview (first 3 codes)")
                            for r in rows[:3]:
                                code  = r.get(st.session_state.cb_code_col, "")
                                label = r.get(st.session_state.cb_label_col, "")
                                st.write(f"- `{code}` — {label}")
                    else:
                        st.warning("Could not read rows from the uploaded file.")
                else:
                    st.session_state.cb_rows    = uploaded_codebook.read()
                    st.session_state.cb_headers = None
            elif not uploaded_codebook and not st.session_state.running:
                st.session_state.cb_rows    = None
                st.session_state.cb_headers = None

        st.divider()

        if st.session_state.running:
            # Stop button — queued click is processed on next rerun after current LLM call
            _doc_noun = "document" if _is_workshop else "interview"
            stop_btn = st.button(
                f"Stop after this {_doc_noun}",
                type="secondary",
                use_container_width=True,
            )
            if stop_btn:
                st.session_state.stop_requested = True
            st.caption("The current LLM call will finish before stopping. All progress is saved.")

        else:
            # In workshop mode the Run button also needs a parsed question list.
            _can_run = bool(uploaded_txts) and (
                bool(st.session_state.questions_list) if _is_workshop else True
            )
            run_btn = st.button(
                "Run Analysis",
                type="primary",
                disabled=not _can_run,
                use_container_width=True,
            )
            if _is_workshop and uploaded_txts and not st.session_state.questions_list:
                st.caption("Add at least one question file to enable Run.")

            # Resume button — only show when a partial run exists
            _resumable = _find_resumable_run()
            if _resumable and not st.session_state.results:
                _n_done  = len(_resumable.get("completed", []))
                _n_total = len(_resumable.get("interviews", []))
                _r_mode  = _resumable.get("mode", "interviews")
                resume_btn = st.button(
                    f"↺  Resume last run  ({_r_mode}, {_n_done} / {_n_total} done)",
                    use_container_width=True,
                )
                if resume_btn:
                    _run_dir  = Path(_resumable["run_dir"])
                    _work_dir = Path(_resumable["work_dir"])
                    _partial  = {}
                    for _fname in _resumable.get("completed", []):
                        _iid = Path(_fname).stem
                        _partial[_iid] = _load_interview_results(_iid, _run_dir)
                    _pending = [
                        str(_work_dir / _f)
                        for _f in _resumable["interviews"]
                        if _f not in _resumable.get("completed", [])
                    ]
                    # Reload the workshop_id map for workshop-mode resumes.
                    _ws_path = _run_dir / "workshops.yaml"
                    _ws_map  = (
                        yaml.safe_load(_ws_path.read_text(encoding="utf-8")) or {}
                        if _ws_path.exists() else {}
                    )

                    st.session_state.run_dir        = _run_dir
                    st.session_state.work_dir       = str(_work_dir)
                    st.session_state.pending        = _pending
                    st.session_state.all_interviews = _resumable["interviews"]
                    st.session_state.codebook_path  = _resumable.get("codebook_path")
                    st.session_state.questions_path = _resumable.get("questions_path")
                    st.session_state.mode           = _r_mode
                    if _resumable.get("anonymise_workshop") is not None:
                        st.session_state.anonymise_workshop = bool(_resumable["anonymise_workshop"])
                    st.session_state.workshop_ids   = _ws_map
                    st.session_state.run_log        = []
                    st.session_state.partial_results = _partial
                    st.session_state.stop_requested = False
                    st.session_state.running        = True
                    st.session_state.results        = None
                    st.rerun()

            # ── Start new run ────────────────────────────────────────────────
            if run_btn and uploaded_txts and _can_run:
                _run_id = datetime.now().strftime("%Y-%m-%d_%H-%M")
                _work   = INSTALL_DIR / "work" / _run_id
                _work.mkdir(parents=True, exist_ok=True)

                _paths = []
                for _f in uploaded_txts:
                    _dest = _work / _f.name
                    _dest.write_bytes(_f.read())
                    _paths.append(_dest)

                _cb_path  = None
                _qs_path  = None

                if _is_workshop:
                    # Persist the parsed questions list — analyse_questions reads it back.
                    _qs_path = _work / "questions.yaml"
                    _qs_path.write_text(
                        yaml.dump(st.session_state.questions_list,
                                  allow_unicode=True, sort_keys=False),
                        encoding="utf-8",
                    )
                else:
                    if st.session_state.cb_rows is not None:
                        _cb_path = _work / "codebook.yaml"
                        if isinstance(st.session_state.cb_rows, (bytes, bytearray)):
                            _cb_path.write_bytes(st.session_state.cb_rows)
                        else:
                            _cb_path.write_text(
                                codebook_rows_to_yaml(
                                    st.session_state.cb_rows,
                                    st.session_state.cb_code_col,
                                    st.session_state.cb_label_col,
                                    st.session_state.cb_desc_col,
                                ),
                                encoding="utf-8",
                            )

                _cfg     = load_cfg(
                    _cb_path,
                    mode=st.session_state.mode,
                    questions_path=_qs_path,
                    anonymise_workshop=(
                        st.session_state.anonymise_workshop if _is_workshop else None
                    ),
                )
                _run_dir = INSTALL_DIR / "output" / _run_id
                (_run_dir / "anonymised").mkdir(parents=True, exist_ok=True)
                (_run_dir / "analysis").mkdir(exist_ok=True)
                (_run_dir / _cfg["gdpr"]["entities_subdir"]).mkdir(exist_ok=True)

                # Workshop mode: assign workshop_NN to each uploaded file and
                # persist the mapping so a resumed run keeps the same IDs.
                _workshop_ids: dict[str, str] = {}
                if _is_workshop:
                    _workshop_ids = _assign_workshop_ids([p.name for p in _paths])
                    (_run_dir / "workshops.yaml").write_text(
                        yaml.dump(_workshop_ids, allow_unicode=True, sort_keys=True),
                        encoding="utf-8",
                    )

                st.session_state.run_dir        = _run_dir
                st.session_state.work_dir       = str(_work)
                st.session_state.pending        = [str(p) for p in _paths]
                st.session_state.all_interviews = [p.name for p in _paths]
                st.session_state.codebook_path  = str(_cb_path) if _cb_path else None
                st.session_state.questions_path = str(_qs_path) if _qs_path else None
                st.session_state.workshop_ids   = _workshop_ids
                st.session_state.run_log        = []
                st.session_state.partial_results = {}
                st.session_state.stop_requested = False
                st.session_state.running        = True
                st.session_state.results        = None
                st.rerun()

    # ── Right column: live status ─────────────────────────────────────────────

    with col_status:
        if st.session_state.running:
            _pending  = st.session_state.pending or []
            _partial  = st.session_state.partial_results or {}
            _all_ivs  = st.session_state.all_interviews or []
            _n_done   = len(_partial)
            _n_total  = len(_all_ivs)
            _doc_noun = "documents" if st.session_state.mode == "workshop" else "interviews"

            st.progress(
                _n_done / _n_total if _n_total else 0.0,
                text=f"**{_n_done} of {_n_total}** {_doc_noun} complete",
            )

            _ws_ids = st.session_state.workshop_ids or {}
            for _fname in _all_ivs:
                _iid = Path(_fname).stem
                _wid = _ws_ids.get(_iid)
                _prefix = f"`{_wid}` &nbsp; " if _wid else ""
                if _iid in _partial:
                    st.markdown(f"✅ &nbsp; {_prefix}**{_iid}**")
                elif _pending and Path(_pending[0]).stem == _iid:
                    st.markdown(f"🔄 &nbsp; {_prefix}**{_iid}** — processing…")
                else:
                    st.markdown(f"⏳ &nbsp; {_prefix}{_iid}")

            st.divider()

            # Honour stop request (queued button click from previous run)
            if st.session_state.stop_requested:
                st.session_state.running        = False
                st.session_state.stop_requested = False
                st.warning(
                    f"Stopped after last completed {_doc_noun[:-1]}. "
                    "Click **↺ Resume last run** in the left column to continue."
                )

            elif _pending:
                # Process the next document
                _path    = Path(_pending[0])
                _iid     = _path.stem
                _cfg     = load_cfg(
                    Path(st.session_state.codebook_path)
                    if st.session_state.codebook_path else None,
                    mode=st.session_state.mode,
                    questions_path=(
                        Path(st.session_state.questions_path)
                        if st.session_state.questions_path else None
                    ),
                    anonymise_workshop=(
                        st.session_state.anonymise_workshop
                        if st.session_state.mode == "workshop" else None
                    ),
                )
                _run_dir = Path(st.session_state.run_dir)

                _doc_log: dict = {"document_id": _iid,
                                  "workshop_id": _ws_ids.get(_iid),
                                  "stages": []}

                with st.status(f"Processing {_iid}…", expanded=True) as _status_box:

                    _stage_bar = st.empty()

                    def _on_stage_start(stage, estimate, _box=_status_box,
                                        _bar=_stage_bar) -> None:
                        label = stage_label(stage)
                        _box.update(label=f"{_iid}: {label}…")
                        _bar.progress(
                            0.0,
                            text=f"{label} — 0%  ·  est. {format_duration(estimate)}",
                        )

                    def _on_stage_tick(stage, elapsed, estimate, tokens, note,
                                       _bar=_stage_bar) -> None:
                        pct   = percent_complete(elapsed, estimate)
                        extra = f" · {note}" if note else ""
                        _bar.progress(
                            pct,
                            text=f"{stage_label(stage)} — {pct * 100:.0f}%  ·  "
                                 f"{tokens:,} tokens{extra}  ·  {elapsed}s",
                        )

                    def _on_stage_done(stage, duration, _log=_doc_log) -> None:
                        st.write(
                            f"✅ &nbsp; **{stage_label(stage)}** — "
                            f"{format_duration(duration)}"
                        )
                        _log["stages"].append({"stage": stage, "duration_s": round(duration, 1)})

                    _result = _process_one_interview(
                        _path, _run_dir, _cfg,
                        _on_stage_start, _on_stage_tick, _on_stage_done,
                    )
                    _stage_bar.empty()
                    # Leave the per-document log expanded so the researcher
                    # can scroll back through stage timings after completion.
                    _status_box.update(
                        state="complete",
                        label=f"✓  {_iid} complete",
                        expanded=True,
                    )

                _doc_log["total_s"] = round(sum(s["duration_s"] for s in _doc_log["stages"]), 1)
                st.session_state.run_log.append(_doc_log)

                st.session_state.partial_results[_iid] = _result
                st.session_state.pending = _pending[1:]
                _write_checkpoint(
                    _run_dir,
                    Path(st.session_state.work_dir),
                    st.session_state.all_interviews,
                    st.session_state.codebook_path,
                    list(st.session_state.partial_results.keys()),
                    mode=st.session_state.mode,
                    questions_path=st.session_state.questions_path,
                    anonymise_workshop=(
                        st.session_state.anonymise_workshop
                        if st.session_state.mode == "workshop" else None
                    ),
                )
                st.rerun()

            else:
                # All documents done — corpus comparison
                _cfg     = load_cfg(
                    Path(st.session_state.codebook_path)
                    if st.session_state.codebook_path else None,
                    mode=st.session_state.mode,
                    questions_path=(
                        Path(st.session_state.questions_path)
                        if st.session_state.questions_path else None
                    ),
                    anonymise_workshop=(
                        st.session_state.anonymise_workshop
                        if st.session_state.mode == "workshop" else None
                    ),
                )
                _run_dir    = Path(st.session_state.run_dir)
                _corpus_dir = _run_dir / "corpus"
                _corpus_dir.mkdir(exist_ok=True)

                with st.status("Building corpus comparison…", expanded=True) as _status_box:

                    _corpus_bar = st.empty()
                    _corpus_est = estimate_seconds(
                        "compare",
                        n_interviews=len(st.session_state.partial_results),
                    )
                    _corpus_t0  = time.time()

                    def _corpus_tick(tokens: int, elapsed: int,
                                     _bar=_corpus_bar, _est=_corpus_est) -> None:
                        pct = percent_complete(elapsed, _est)
                        _bar.progress(
                            pct,
                            text=f"Corpus comparison — {pct * 100:.0f}%  ·  "
                                 f"{tokens:,} tokens  ·  {elapsed}s",
                        )

                    _sf       = sorted((_run_dir / "analysis").glob("*_summary.json"))
                    _is_ws_run = _cfg.get("mode") == "workshop"
                    _stage3_glob = "*_questions.json" if _is_ws_run else "*_themes.json"
                    _tf       = sorted((_run_dir / "analysis").glob(_stage3_glob))

                    # Workshop mode: look for filename-similar groups before
                    # asking the LLM to synthesise. The LLM gets the clusters
                    # as a hint so it can verify them against the summaries.
                    _potential_dupes = []
                    if _is_ws_run:
                        _potential_dupes = detect_workshop_duplicates(
                            st.session_state.all_interviews,
                            st.session_state.workshop_ids or {},
                        )
                        (_corpus_dir / "potential_duplicates.json").write_text(
                            json.dumps(_potential_dupes, ensure_ascii=False, indent=2),
                            encoding="utf-8",
                        )

                    _corpus   = build_corpus_comparison(
                        _sf, _tf, _cfg,
                        tick_cb=_corpus_tick,
                        potential_duplicates=_potential_dupes,
                        workshop_ids=(
                            st.session_state.workshop_ids if _is_ws_run else None
                        ),
                    )
                    _corpus["_duration"] = time.time() - _corpus_t0
                    _corpus["_potential_duplicates"] = _potential_dupes
                    _matrix_name = (
                        "questions_matrix.json" if _is_ws_run else "themes_matrix.json"
                    )
                    (_corpus_dir / _matrix_name).write_text(
                        json.dumps(_corpus["matrix"], ensure_ascii=False, indent=2),
                        encoding="utf-8",
                    )
                    (_corpus_dir / "comparison_report.md").write_text(
                        _corpus["report"], encoding="utf-8"
                    )
                    _corpus_bar.empty()
                    st.write(
                        f"✅ &nbsp; **Corpus comparison** — "
                        f"{format_duration(_corpus['_duration'])}"
                    )
                    _status_box.update(
                        state="complete",
                        label=f"Corpus comparison — "
                              f"{format_duration(_corpus['_duration'])}",
                        expanded=True,
                    )

                st.session_state.results = {
                    **st.session_state.partial_results,
                    "_corpus":       _corpus,
                    "_mode":         _cfg.get("mode") or "interviews",
                    "_workshop_ids": st.session_state.workshop_ids or {},
                    "_run_log":      list(st.session_state.run_log),
                }
                st.session_state.running = False
                # Remove checkpoint — run fully complete
                _cp = _run_dir / "checkpoint.json"
                if _cp.exists():
                    _cp.unlink()
                st.rerun()

        elif st.session_state.results:
            st.success("Analysis complete. Switch to the **Outcomes** tab to view results.")
            st.caption(f"Output: `{st.session_state.run_dir}`")

            # Processing log: stays visible after the job so researchers can
            # scroll back through per-stage timings.
            _r        = st.session_state.results
            _is_ws    = _r.get("_mode") == "workshop"
            _wsmap    = _r.get("_workshop_ids", {}) or {}
            _run_log  = _r.get("_run_log") or []
            _corpus_d = (_r.get("_corpus") or {}).get("_duration")
            _dupes    = (_r.get("_corpus") or {}).get("_potential_duplicates") or []

            with st.expander("Processing log", expanded=True):
                if _is_ws and _dupes:
                    st.markdown(
                        f"**Potential duplicate workshops detected:** {len(_dupes)} cluster(s). "
                        "See the executive summary for details."
                    )
                if _run_log:
                    for rec in _run_log:
                        _doc_id = rec.get("document_id", "")
                        _wid    = rec.get("workshop_id") or _wsmap.get(_doc_id) or ""
                        _prefix = f"`{_wid}` &nbsp; " if _wid else ""
                        st.markdown(
                            f"**{_prefix}{_doc_id}** &nbsp; — &nbsp; "
                            f"{format_duration(rec.get('total_s', 0))}"
                        )
                        _stage_bits = "  ·  ".join(
                            f"{stage_label(s['stage'])} {format_duration(s['duration_s'])}"
                            for s in rec.get("stages", [])
                        )
                        if _stage_bits:
                            st.caption(_stage_bits)
                else:
                    st.caption("No per-document log captured (resumed run).")

                if _corpus_d:
                    st.markdown(
                        f"**Corpus comparison** &nbsp; — &nbsp; "
                        f"{format_duration(_corpus_d)}"
                    )

        else:
            if st.session_state.mode == "workshop":
                st.info(
                    "Upload workshop description sheets and at least one question "
                    "file on the left, then click **Run Analysis** to begin.\n\n"
                    "Progress is saved after each document — you can stop at any time and resume later."
                )
            else:
                st.info(
                    "Upload transcript files on the left and click **Run Analysis** to begin.\n\n"
                    "Progress is saved after each interview — you can stop at any time and resume later."
                )

# ── Outcomes tab ──────────────────────────────────────────────────────────────

with tab_outcomes:
    if not st.session_state.results:
        st.info("Run an analysis first — or load a previous run from disk below.")

        _loadable = _list_loadable_runs()
        if _loadable:
            _options = [
                f"{p.name}  ·  "
                f"{'workshop' if (p / 'corpus' / 'questions_matrix.json').exists() else 'interviews'}"
                f"  ·  {len(list((p / 'analysis').glob('*_summary.json')))} doc(s)"
                for p in _loadable
            ]
            _sel_col, _btn_col = st.columns([4, 1], gap="medium")
            with _sel_col:
                _picked = st.selectbox(
                    "Load a previous run",
                    options=list(range(len(_options))),
                    format_func=lambda i: _options[i],
                    key="load_run_selectbox",
                )
            with _btn_col:
                st.write("")  # vertical alignment with selectbox
                _load_btn = st.button(
                    "Load",
                    type="primary",
                    use_container_width=True,
                    key="load_run_button",
                )
            if _load_btn and _picked is not None:
                _chosen = _loadable[_picked]
                _r = _load_complete_run(_chosen)
                st.session_state.results = _r
                st.session_state.run_dir = _chosen
                st.session_state.mode    = _r.get("_mode") or "interviews"
                st.rerun()
        else:
            st.caption("No completed runs found under `output/`.")
    else:
        results  = st.session_state.results
        run_dir: Path = Path(st.session_state.run_dir)

        _file_slug = (
            "workshop-analysis" if results.get("_mode") == "workshop"
            else "interview-analysis"
        )
        dl_col1, dl_col2 = st.columns(2, gap="large")
        with dl_col1:
            st.download_button(
                "Download Word report (.docx)",
                data=generate_docx_report(run_dir, results),
                file_name=f"{_file_slug}-{run_dir.name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
            )
        with dl_col2:
            st.download_button(
                "Download HTML report",
                data=generate_html_report(run_dir, results).encode("utf-8"),
                file_name=f"{_file_slug}-{run_dir.name}.html",
                mime="text/html",
                use_container_width=True,
            )
        st.caption(
            "Open the HTML report in any browser. "
            "Use File › Print › Save as PDF for a PDF copy."
        )

        st.subheader("Report highlights")
        st.markdown(
            f'<div class="riecs-scroll-pane">'
            f'{_report_highlights_html(results)}'
            f'</div>',
            unsafe_allow_html=True,
        )
